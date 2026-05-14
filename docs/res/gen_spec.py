#!/usr/bin/env python3
"""Generate AssetCore_Technical_Specification.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = "/home/miyano/frappe-bench/apps/assetcore/docs/res/AssetCore_Technical_Specification.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.5)

# ── Styles helpers ────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def note(text):
    p = doc.add_paragraph(text)
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    p.runs[0].font.size = Pt(10)
    return p

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2E74B5')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        c._tc.get_or_add_tcPr().append(shd)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri, row in enumerate(rows):
        trow = t.rows[ri+1]
        for ci, val in enumerate(row):
            c = trow.cells[ci]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'EBF3FB')
                shd.set(qn('w:val'), 'clear')
                c._tc.get_or_add_tcPr().append(shd)
    return t

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TRANG BÌA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ASSETCORE")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Hệ thống Quản lý Vòng đời Thiết bị Y tế")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\nTÀI LIỆU THUYẾT MINH KỸ THUẬT PHẦN MỀM")
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Technical Specification Document")
run.font.size = Pt(13)
run.italic = True

doc.add_paragraph("\n\n")

info_rows = [
    ("Phiên bản", "1.0.0"),
    ("Ngày ban hành", "2026-05-13"),
    ("Nền tảng", "Frappe Framework v15 / Python 3.x / MariaDB"),
    ("Chuẩn tham chiếu", "WHO HTM 2025, NĐ 98/2021, ISO 13485:2016, ISO/IEC 17025"),
    ("Trạng thái", "Đang triển khai — Wave 1 Live, Wave 2 Live, Wave 3 Planned"),
]
t = doc.add_table(rows=len(info_rows), cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info_rows):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    t.rows[i].cells[1].text = v

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# MỤC LỤC (thủ công)
# ══════════════════════════════════════════════════════════════════════════════
h1("MỤC LỤC")
toc = [
    "1. Tổng quan hệ thống",
    "   1.1 Giới thiệu AssetCore",
    "   1.2 Bài toán giải quyết",
    "   1.3 Kiến trúc tổng thể",
    "   1.4 WHO HTM Lifecycle",
    "   1.5 Nguyên tắc kiến trúc",
    "2. Kiến trúc module",
    "   2.1 Phân nhóm 4 khối",
    "   2.2 Danh sách module IMM-00 đến IMM-17",
    "   2.3 Trạng thái triển khai",
    "3. Mô tả chức năng từng module",
    "   3.0 IMM-00 — Foundation / Master Cross-cutting",
    "   3.1 IMM-01 — Needs Assessment & Budget Estimation",
    "   3.2 IMM-02 — Tech Spec & Market Analysis",
    "   3.3 IMM-03 — Vendor Evaluation & Procurement",
    "   3.4 IMM-04 — Lắp đặt, Định danh & Kiểm tra Ban đầu",
    "   3.5 IMM-05 — Hồ sơ thiết bị (Asset Documents)",
    "   3.6 IMM-06 — Đào tạo & Năng lực",
    "   3.7 IMM-08 — Bảo trì định kỳ (PM)",
    "   3.8 IMM-09 — Sửa chữa (Corrective Maintenance)",
    "   3.9 IMM-11 — Hiệu chuẩn (Calibration)",
    "   3.10 IMM-12 — Sự cố & CAPA (Incident & RCA)",
    "   3.11 IMM-15 — Spare Parts Inventory Tracking",
    "   3.12 IMM-16 — Giám sát Tuân thủ & CAPA",
    "   3.13 IMM-07/10/13/14/17 — Nhóm Wave 3",
    "4. Data model tổng thể",
    "   4.1 Master Data",
    "   4.2 Operational Data",
    "   4.3 Governance Data",
    "   4.4 Danh mục DocType đầy đủ",
    "5. Tích hợp hệ thống ngoài",
    "6. Workflow engine & SLA",
    "   6.1 Work Order Engine",
    "   6.2 SLA Policy",
    "   6.3 Lifecycle Event",
    "   6.4 Audit Trail (SHA-256 Chain)",
    "7. Bảo mật & Phân quyền",
    "8. Yêu cầu phi chức năng",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.runs[0].font.size = Pt(11)
    if not item.startswith(" "):
        p.runs[0].bold = True

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 1 — TỔNG QUAN HỆ THỐNG
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Tổng quan hệ thống")

h2("1.1 Giới thiệu AssetCore")
para(
    "AssetCore là hệ thống quản lý vòng đời thiết bị y tế (Health Technology Management — HTM) "
    "xây dựng trên nền tảng Frappe Framework v15. Hệ thống được thiết kế đặc thù cho các cơ sở y tế "
    "Việt Nam, tuân thủ chuẩn WHO HTM 2025, Nghị định 98/2021/NĐ-CP về quản lý trang thiết bị y tế, "
    "ISO 13485:2016 và ISO/IEC 17025."
)
para(
    "AssetCore KHÔNG phải CMMS (Computerized Maintenance Management System) đơn lẻ mà là một "
    "operating architecture quản trị toàn bộ vòng đời thiết bị y tế — từ xác định nhu cầu, mua sắm, "
    "lắp đặt, vận hành, bảo trì, hiệu chuẩn cho đến giải nhiệm — với audit trail bất biến và "
    "compliance tự động theo quy định pháp luật."
)

h2("1.2 Bài toán giải quyết")
para(
    "Tại phần lớn bệnh viện công Việt Nam, quản lý thiết bị y tế đang đối mặt với các vấn đề:"
)
bullet("Không có audit trail — không biết ai phê duyệt gì, khi nào")
bullet("Workflow ngầm, không có engine, không SLA — không đo được KPI tuân thủ")
bullet("Lịch PM/hiệu chuẩn quản lý bằng Excel/sổ giấy, dễ bỏ sót")
bullet("Incident Report qua điện thoại/Zalo, không có biểu mẫu chuẩn, không RCA")
bullet("Thiết bị fail hiệu chuẩn vẫn tiếp tục sử dụng trên bệnh nhân")
bullet("Không liên kết sang hệ thống tài chính/kế toán tài sản")
bullet("Module rời rạc, không lifecycle — không truy vết được lịch sử đầy đủ")

h2("1.3 Kiến trúc tổng thể")
para(
    "AssetCore được xây dựng theo kiến trúc 7 lớp, từ lớp người dùng xuống lớp dữ liệu và tích hợp:"
)
add_table(
    ["Lớp kiến trúc", "Mô tả", "Công nghệ"],
    [
        ("Lớp người dùng", "Vue 3 + TypeScript + Pinia + TailwindCSS", "Frontend SPA"),
        ("Lớp workflow & SLA", "Frappe Workflow Engine + SLA Policy", "17 workflow, 22 states máy trạng thái"),
        ("Lớp nghiệp vụ (IMM)", "17 module IMM-00 → IMM-17, chia 4 khối", "Python service layer 3-tier"),
        ("Lớp dữ liệu", "DocType ORM → MariaDB", "70+ DocType, audit trail SHA-256"),
        ("Lớp tích hợp", "REST/OpenAPI + FHIR-ready", "Frappe Whitelist API"),
        ("Lớp phân tích", "Dashboard KPI, drill-down về bản ghi nguồn", "Frappe Reports + Vue"),
        ("Lớp QMS & governance", "Audit Trail, CAPA, Internal Audit", "IMM-00 foundation"),
    ]
)

h2("1.4 WHO HTM Lifecycle")
para(
    "AssetCore ánh xạ toàn bộ 6 phase của WHO HTM lifecycle sang các module nghiệp vụ:"
)
add_table(
    ["Phase WHO HTM", "Module AssetCore", "Mô tả"],
    [
        ("Needs Assessment", "IMM-01", "Xác định nhu cầu, ưu tiên, lập kế hoạch ngân sách"),
        ("Procurement", "IMM-02, IMM-03", "Thông số kỹ thuật, đánh giá NCC, quyết định mua sắm"),
        ("Installation & Commissioning", "IMM-04, IMM-05, IMM-06", "Lắp đặt, định danh, hồ sơ, đào tạo"),
        ("Operation", "IMM-00 (AC Asset), IMM-07", "Registry thiết bị, theo dõi hiệu suất"),
        ("Maintenance", "IMM-08, IMM-09, IMM-11, IMM-12, IMM-15", "PM, CM, Calibration, Incident, Spare Parts"),
        ("Decommission", "IMM-13, IMM-14", "Ngừng sử dụng, điều chuyển, giải nhiệm"),
    ]
)

h2("1.5 Nguyên tắc kiến trúc")
add_table(
    ["Nguyên tắc", "Ràng buộc"],
    [
        ("Chỉ phụ thuộc Frappe v15", "KHÔNG cần ERPNext — tất cả DocType prefix AC/IMM được tái tạo native"),
        ("Tách đúng domain", "Item ≠ Model ≠ Asset ≠ Event ≠ Work Order"),
        ("Audit trail bắt buộc", "Mọi state mutation → log_audit_event() với SHA-256 chain"),
        ("Workflow trước UI", "KHÔNG build UI trước khi có workflow + SLA"),
        ("3-tier strict", "API layer → Service layer → Repository/ORM — không logic trong controller"),
        ("TDD bắt buộc", "Viết test trước implement — CLAUDE.md §17"),
        ("Không bypass audit", "transition_asset_status() và log_audit_event() không được bypass"),
    ]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 2 — KIẾN TRÚC MODULE
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Kiến trúc module")

h2("2.1 Phân nhóm 4 khối")
para(
    "17 module IMM được tổ chức thành 4 khối chức năng theo WHO HTM lifecycle:"
)
add_table(
    ["Khối", "Tên", "Module", "Mô tả"],
    [
        ("A — Khối 1", "Planning & Procurement", "IMM-01, 02, 03", "Xác định nhu cầu → thông số kỹ thuật → đánh giá NCC → mua sắm"),
        ("B — Khối 2", "Deployment", "IMM-04, 05, 06", "Lắp đặt → hồ sơ thiết bị → đào tạo vận hành"),
        ("C — Khối 3", "Operations & Maintenance", "IMM-07, 08, 09, 10, 11, 12, 15, 16, 17", "Vận hành → PM → CM → Calibration → Incident/CAPA → Spare Parts → Compliance"),
        ("D — Khối 4", "End-of-Life", "IMM-13, 14", "Ngừng sử dụng/điều chuyển → giải nhiệm thiết bị"),
    ]
)
para(
    "Ngoài 17 module nghiệp vụ, IMM-00 là foundation cross-cutting layer cung cấp:"
)
bullet("DocType lõi (AC Asset, AC Supplier, AC Location, AC Department, AC Asset Category)")
bullet("Shared services (transition_asset_status, log_audit_event, create_lifecycle_event)")
bullet("8 vai trò RBAC (IMM System Admin, Department Head, Operations Manager, Workshop Lead, Technician, QA Officer, Document Officer, Reporting User)")
bullet("SLA Policy engine (giải quyết SLA theo risk_class × priority × asset_category)")
bullet("Scheduler jobs (4 daily jobs: SLA breach, PM overdue, calibration alert, CAPA escalation)")

h2("2.2 Danh sách module IMM-00 đến IMM-17")
add_table(
    ["Module", "Tên đầy đủ", "Khối", "Wave", "Trạng thái"],
    [
        ("IMM-00", "Foundation / Master Cross-cutting", "Cross-cutting", "Master", "Live"),
        ("IMM-01", "Needs Assessment & Budget Estimation", "A — Khối 1", "Wave 2", "Live"),
        ("IMM-02", "Tech Spec & Market Analysis", "A — Khối 1", "Wave 2", "Live"),
        ("IMM-03", "Vendor Evaluation & Procurement Decision", "A — Khối 1", "Wave 2", "Live"),
        ("IMM-04", "Lắp đặt, Định danh & Kiểm tra Ban đầu", "B — Khối 2", "Wave 1", "Live"),
        ("IMM-05", "Hồ sơ thiết bị (Asset Documents)", "B — Khối 2", "Wave 1", "Live"),
        ("IMM-06", "Đào tạo & Năng lực (Training & Competency)", "B — Khối 2", "Wave 2", "Live"),
        ("IMM-07", "Theo dõi hiệu suất (Performance Tracking)", "C — Khối 3", "Wave 3", "Planned"),
        ("IMM-08", "Bảo trì định kỳ (Preventive Maintenance)", "C — Khối 3", "Wave 1", "Live"),
        ("IMM-09", "Sửa chữa (Corrective Maintenance)", "C — Khối 3", "Wave 1", "Live"),
        ("IMM-10", "Hậu kiểm và tuân thủ (Post-market Surveillance)", "C — Khối 3", "Wave 3", "Planned"),
        ("IMM-11", "Hiệu chuẩn (Calibration)", "C — Khối 3", "Wave 1", "Live"),
        ("IMM-12", "Sự cố & CAPA (Incident, RCA, CAPA)", "C — Khối 3", "Wave 1", "Live"),
        ("IMM-13", "Ngừng sử dụng và điều chuyển (Decommission & Transfer)", "D — Khối 4", "Wave 3", "Planned"),
        ("IMM-14", "Giải nhiệm thiết bị (Formal Retirement)", "D — Khối 4", "Wave 3", "Planned"),
        ("IMM-15", "Spare Parts Inventory Tracking", "C — Khối 3", "Wave 3", "Planned"),
        ("IMM-16", "Giám sát Tuân thủ & CAPA (Compliance Monitoring)", "C — Khối 3", "Wave 3", "Planned"),
        ("IMM-17", "Phân tích dự đoán (Predictive Analytics)", "C — Khối 3", "Wave 3", "Planned"),
    ]
)

h2("2.3 Trạng thái triển khai")
add_table(
    ["Wave", "Scope", "Module", "Trạng thái"],
    [
        ("Wave 1", "6 module vận hành cốt lõi", "IMM-04, 05, 08, 09, 11, 12", "Live — BE + FE deployed"),
        ("Wave 2", "4 module lập kế hoạch & triển khai", "IMM-01, 02, 03, 06", "Live — BE + FE deployed"),
        ("Wave 3", "7 module nâng cao", "IMM-07, 10, 13, 14, 15, 16, 17", "Planned — Docs có, chưa implement"),
    ]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 3 — MÔ TẢ CHỨC NĂNG TỪNG MODULE
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Mô tả chức năng từng module")

# ── IMM-00 ───────────────────────────────────────────────────────────────────
h2("3.0 IMM-00 — Foundation / Master Cross-cutting")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên", "Foundation / Master Cross-cutting"),
        ("Khối", "Cross-cutting — nền tảng cho tất cả module"),
        ("Wave", "Master — deploy trước mọi module"),
        ("Actor chính", "IMM System Admin, IMM Operations Manager"),
        ("Chuẩn tham chiếu", "WHO HTM 2025, NĐ 98/2021, ISO 13485:2016"),
    ]
)
h3("Mục đích")
para(
    "IMM-00 là foundation layer tự chứa. Không có workflow nghiệp vụ riêng mà cung cấp "
    "toàn bộ hạ tầng dùng chung: DocType lõi, shared services, vai trò RBAC, SLA engine, "
    "audit trail và scheduler jobs cho toàn hệ thống."
)
h3("DocType chính")
add_table(
    ["DocType", "Prefix", "Mô tả"],
    [
        ("AC Asset", "AC-", "Registry thiết bị vật lý — trục trung tâm của toàn hệ thống"),
        ("AC Supplier", "AC-SUP-", "Nhà cung cấp — có thể là vendor, lab hiệu chuẩn, NCC phụ tùng"),
        ("AC Location", "AC-LOC-", "Vị trí vật lý: tòa nhà → tầng → phòng → khu vực"),
        ("AC Department", "AC-DEPT-", "Khoa/phòng bệnh viện"),
        ("AC Asset Category", "—", "Phân loại thiết bị (theo IFU/risk class)"),
        ("IMM Device Model", "IMM-MDL-", "Catalog model thiết bị: thông số kỹ thuật, PM interval, cal interval"),
        ("IMM SLA Policy", "—", "Chính sách SLA theo risk_class × priority × asset_category"),
        ("IMM Audit Trail", "IMM-AUD-", "Audit trail bất biến — SHA-256 chain"),
        ("IMM CAPA Record", "IMM-CAP-", "Corrective and Preventive Action"),
        ("Asset Lifecycle Event", "—", "Sự kiện vòng đời thiết bị — immutable"),
        ("Incident Report", "IR-", "Báo cáo sự cố thiết bị (dùng chung cho IMM-12)"),
    ]
)
h3("Shared Services (Python)")
add_table(
    ["Function", "Mô tả"],
    [
        ("transition_asset_status()", "Chuyển trạng thái lifecycle của AC Asset — bắt buộc dùng, không bypass"),
        ("log_audit_event()", "Ghi audit trail với SHA-256 chain — bắt buộc cho mọi state mutation"),
        ("create_lifecycle_event()", "Tạo Asset Lifecycle Event row immutable"),
        ("get_sla_policy() / resolve_sla_policy()", "Giải quyết SLA theo risk_class × priority"),
        ("open_capa() / close_capa_record()", "Tạo/đóng CAPA Record"),
        ("validate_asset_for_operations()", "Kiểm tra asset đủ điều kiện vận hành trước action"),
        ("normalize_filters()", "Chuẩn hóa filter từ FE trước khi query"),
    ]
)
h3("Lifecycle Status Machine của AC Asset")
add_table(
    ["Trạng thái", "Mô tả"],
    [
        ("Needs Assessment", "Mới tạo — chờ phê duyệt nhu cầu"),
        ("Procurement", "Đang mua sắm"),
        ("In Storage", "Nhận hàng, chưa lắp đặt"),
        ("Under Installation", "Đang trong quy trình IMM-04"),
        ("Commissioned", "Đã lắp đặt, chưa bàn giao lâm sàng"),
        ("Active", "Đang vận hành bình thường"),
        ("Under Maintenance", "Đang bảo trì PM"),
        ("Under Repair", "Đang sửa chữa CM"),
        ("Calibrating", "Đang hiệu chuẩn"),
        ("Clinical Hold", "Tạm dừng sử dụng lâm sàng"),
        ("Out of Service", "Ngừng hoạt động tạm thời"),
        ("Decommissioned", "Đã giải nhiệm vĩnh viễn"),
    ]
)

doc.add_paragraph()

# ── IMM-01 ───────────────────────────────────────────────────────────────────
h2("3.1 IMM-01 — Needs Assessment & Budget Estimation")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Nhu cầu thiết bị y tế & Ước tính ngân sách"),
        ("Khối", "A — Khối 1 (Planning & Procurement)"),
        ("Wave", "Wave 2 — Live"),
        ("Actor chính", "IMM Operations Manager, IMM Department Head, IMM System Admin"),
        ("DocType chính", "IMM Needs Request, Budget Estimate Line, Needs Priority Scoring"),
        ("Workflow", "IMM-01 Needs Workflow, IMM-01 Plan Workflow"),
    ]
)
h3("Mục đích")
para(
    "IMM-01 hệ thống hóa việc tiếp nhận và phê duyệt nhu cầu mua sắm thiết bị y tế theo chu kỳ "
    "kế hoạch. Thay thế quy trình xin-cho qua email/giấy tờ bằng digital workflow có phân quyền, "
    "ưu tiên hóa nhu cầu theo tiêu chí (risk class, clinical urgency, budget ceiling) và tích hợp "
    "sang IMM-02 để lập thông số kỹ thuật."
)
h3("Luồng nghiệp vụ")
para("Draft → Submitted → Under Review → Approved / Rejected → (nếu Approved) → Linked to IMM-02")
h3("Tích hợp")
add_table(
    ["Module", "Chiều", "Mô tả"],
    [
        ("IMM-02", "OUTPUT", "Needs Request được approved → trigger tạo Tech Spec"),
        ("IMM-03", "OUTPUT", "Approved Needs → Procurement Plan"),
        ("IMM-00 AC Asset Category", "INPUT", "Phân loại thiết bị cho ưu tiên"),
    ]
)

doc.add_paragraph()

# ── IMM-02 ───────────────────────────────────────────────────────────────────
h2("3.2 IMM-02 — Tech Spec & Market Analysis")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Thông số Kỹ thuật & Phân tích Thị trường"),
        ("Khối", "A — Khối 1"),
        ("Wave", "Wave 2 — Live"),
        ("Actor chính", "IMM Operations Manager, IMM Workshop Lead"),
        ("DocType chính", "IMM Tech Spec, Tech Spec Requirement, Tech Spec Document, IMM Market Benchmark, Benchmark Candidate"),
        ("Workflow", "IMM-02 Spec Workflow"),
    ]
)
h3("Mục đích")
para(
    "IMM-02 số hóa quy trình soạn thảo thông số kỹ thuật (TSKT) cho thiết bị y tế trước khi "
    "đưa ra đấu thầu. Module hỗ trợ lập TSKT từ Device Model catalog, benchmark thị trường "
    "qua so sánh nhà cung cấp, và xuất bản TSKT sang IMM-03 để làm cơ sở đánh giá NCC."
)
h3("Luồng nghiệp vụ")
para("Draft → Under Review → Approved → Published → (linked to IMM-03 Vendor Evaluation)")
h3("Chức năng chính")
bullet("Soạn TSKT theo template chuẩn hóa với các tiêu chí kỹ thuật bắt buộc")
bullet("Benchmark thị trường: so sánh ≥3 nhà cung cấp theo giá, năng lực, chứng chỉ")
bullet("Lock-in risk assessment (IMM Lock-in Risk Assessment DocType)")
bullet("Lưu trữ tài liệu TSKT (IMM AVL Entry cho Approved Vendor List)")

doc.add_paragraph()

# ── IMM-03 ───────────────────────────────────────────────────────────────────
h2("3.3 IMM-03 — Vendor Evaluation & Procurement Decision")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Đánh giá Nhà cung cấp & Quyết định Mua sắm"),
        ("Khối", "A — Khối 1"),
        ("Wave", "Wave 2 — Live"),
        ("Actor chính", "IMM Operations Manager, IMM Department Head, IMM System Admin"),
        ("DocType chính", "IMM Vendor Evaluation, Vendor Eval Candidate, Vendor Eval Criterion, IMM Vendor Scorecard, IMM Procurement Decision, IMM Procurement Plan, Procurement Plan Line"),
        ("Workflow", "IMM-03 AVL Workflow, IMM-03 Decision Workflow, IMM-03 Vendor Eval Workflow"),
    ]
)
h3("Mục đích")
para(
    "IMM-03 chuẩn hóa quy trình đánh giá nhà cung cấp thiết bị y tế theo bộ tiêu chí định lượng "
    "(kỹ thuật, tài chính, dịch vụ hậu mãi, tuân thủ pháp lý) và lập quyết định mua sắm với "
    "phê duyệt đa cấp. Module bảo đảm mọi quyết định mua sắm đều có hồ sơ traceable từ "
    "nhu cầu (IMM-01) đến TSKT (IMM-02) đến NCC được chọn."
)
h3("Luồng nghiệp vụ — Vendor Evaluation")
para("Draft → Submitted → Scoring → Recommended → Approved → AVL Update")
h3("Luồng nghiệp vụ — Procurement Decision")
para("Draft → Under Review → Approved → (PO issued) → Linked to IMM-04")
h3("KPI")
add_table(
    ["KPI", "Mô tả"],
    [
        ("Vendor Score Distribution", "Phân bổ điểm NCC theo tiêu chí"),
        ("Time-to-Decision", "Thời gian từ submit đến decision approved"),
        ("AVL Coverage", "% thiết bị có NCC trong Approved Vendor List"),
    ]
)

doc.add_paragraph()

# ── IMM-04 ───────────────────────────────────────────────────────────────────
h2("3.4 IMM-04 — Lắp đặt, Định danh & Kiểm tra Ban đầu")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Lắp đặt, Định danh & Kiểm tra Ban đầu (Installation & Commissioning)"),
        ("Khối", "B — Khối 2 (Deployment)"),
        ("Wave", "Wave 1 — Live"),
        ("Actor chính", "HTM Technician, Biomed Engineer, Vendor Engineer, QA Officer, Workshop Head, VP Block2"),
        ("DocType chính", "Asset Commissioning (submittable), Commissioning Checklist, Commissioning Document Record, Asset QA Non Conformance"),
        ("Workflow", "IMM-04 Workflow — 11 states, 22 transitions, 6 Gates (G01–G06)"),
        ("Naming series", "ACC-.YY.-.MM.-.#####"),
        ("Số endpoints", "17 REST endpoints"),
    ]
)
h3("Mục đích")
para(
    "IMM-04 là deployment gateway bắt buộc — mọi thiết bị từ khi nhận từ nhà cung cấp đều phải "
    "qua pipeline lắp đặt 11 bước. Không có phiếu IMM-04 ở trạng thái Clinical Release thì "
    "thiết bị không được sử dụng và không có AC Asset record — đảm bảo 100% traceability."
)
h3("Workflow states (11 trạng thái)")
add_table(
    ["Trạng thái", "Mô tả", "Gate"],
    [
        ("Draft", "Phiếu mới tạo, chưa submit", "—"),
        ("Pending Inspection", "Đang kiểm tra hồ sơ nhận hàng", "G01: CO/CQ đủ"),
        ("Installation In Progress", "Đang lắp đặt vật lý", "—"),
        ("Initial Inspection", "Đang đo kiểm an toàn điện và chức năng", "G02: IMM-05 docs đủ"),
        ("QA Review", "QA Officer review với thiết bị risk class C/D", "G03: QA sign-off"),
        ("Clinical Hold", "Tạm dừng — phát sinh non-conformance nghiêm trọng", "—"),
        ("Pending BGD Approval", "Chờ VP Block2 phê duyệt cuối", "G04: BGD signature"),
        ("Clinical Release", "Đã được phép sử dụng lâm sàng — AC Asset được tạo", "G05/G06"),
        ("Cancelled", "Huỷ phiếu", "—"),
        ("DOA — Return", "Dead on Arrival — trả nhà cung cấp", "—"),
        ("On Hold", "Tạm dừng chờ phụ kiện / giấy phép", "—"),
    ]
)
h3("Validation Rules bắt buộc")
add_table(
    ["Rule", "Mô tả"],
    [
        ("VR-01", "Serial number vendor_serial_no phải unique trên toàn hệ thống"),
        ("VR-02", "Thiết bị bức xạ (custom_is_radiation=1) bắt buộc QA sign-off trước Clinical Release"),
        ("VR-03", "Tất cả non-conformance phải được Close hoặc Accept trước submit lên BGD"),
        ("VR-04", "GW-2: IMM-05 document set phải đủ mandatory documents"),
        ("VR-05", "Electrical safety test phải có kết quả Pass trước Initial Inspection completed"),
        ("VR-06", "Lifecycle Event được tạo khi Submit — immutable, không xoá được"),
        ("VR-07", "Cancel chỉ cho phép ở trạng thái Draft/On Hold"),
    ]
)
h3("Output khi Submit (Clinical Release)")
bullet("Tạo AC Asset record với custom_vendor_serial, custom_internal_qr, custom_comm_ref")
bullet("Auto-import document set sang IMM-05 (create_initial_document_set)")
bullet("Fire lifecycle event 'commissioned' → IMM-08 PM Schedule (listener)")
bullet("Publish realtime event imm04_asset_released")
bullet("Tạo Depreciation Schedule (services/depreciation)")
h3("KPI")
add_table(
    ["KPI", "Target"],
    [
        ("Installation cycle time", "≤ 5 ngày làm việc"),
        ("Gate compliance rate", "100% — 0 thiết bị qua Clinical Release khi thiếu gate"),
        ("Serial uniqueness", "0 duplicate"),
        ("QA review rate cho Class C/D", "100%"),
    ]
)

doc.add_paragraph()

# ── IMM-05 ───────────────────────────────────────────────────────────────────
h2("3.5 IMM-05 — Hồ sơ thiết bị (Asset Documents)")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Hồ sơ thiết bị — Quản lý tài liệu kỹ thuật"),
        ("Khối", "B — Khối 2"),
        ("Wave", "Wave 1 — Live"),
        ("Actor chính", "HTM Technician, Biomed Engineer, IMM Document Officer"),
        ("DocType chính", "Asset Document (submittable), Required Document Type, Document Request"),
        ("Workflow", "IMM-05 Document Workflow"),
        ("Số endpoints", "14 REST endpoints"),
    ]
)
h3("Mục đích")
para(
    "IMM-05 quản lý toàn bộ hồ sơ kỹ thuật của thiết bị y tế theo vòng đời: CO (Certificate of Origin), "
    "CQ (Certificate of Quality), IFU (Instructions for Use), Service Manual, Warranty Certificate, "
    "Radiation Permit và các tài liệu bắt buộc theo NĐ 98/2021. Module đảm bảo completeness score "
    "và cảnh báo tài liệu sắp hết hạn."
)
h3("Chức năng chính")
bullet("Tiếp nhận tài liệu tự động từ IMM-04 commissioning (create_initial_document_set)")
bullet("Quản lý version tài liệu — lưu trữ nhiều phiên bản, archive version cũ")
bullet("Completeness score theo required_document_type profile của asset category")
bullet("Expiry alert scheduler: cảnh báo trước 30/60/90 ngày")
bullet("Document request workflow: yêu cầu bổ sung tài liệu còn thiếu")
bullet("Read-only public link cho audit/kiểm tra ngoài — không cấp quyền edit")
h3("Tích hợp")
add_table(
    ["Module", "Chiều", "Mô tả"],
    [
        ("IMM-04", "INPUT", "Auto-populate document set khi commissioning"),
        ("IMM-04 Gate GW-2", "GATE", "IMM-04 kiểm tra completeness qua IMM-05 trước Clinical Release"),
        ("IMM-09", "INPUT", "Post-repair tài liệu cập nhật (firmware, service report)"),
        ("IMM-11", "INPUT", "Calibration certificate lưu trữ"),
    ]
)

doc.add_paragraph()

# ── IMM-06 ───────────────────────────────────────────────────────────────────
h2("3.6 IMM-06 — Đào tạo & Năng lực (Training & Competency)")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Đào tạo & Năng lực vận hành thiết bị y tế"),
        ("Khối", "B — Khối 2"),
        ("Wave", "Wave 2 — Live"),
        ("Actor chính", "IMM Workshop Lead, IMM Technician, IMM Operations Manager"),
        ("DocType chính", "IMM Training Program, IMM Training Session, IMM Training Participant, IMM User Competency, IMM Competency Gap Report"),
        ("Workflow", "IMM-06 Session Workflow, IMM-06 Competency Workflow"),
    ]
)
h3("Mục đích")
para(
    "IMM-06 quản lý chương trình đào tạo và hồ sơ năng lực của đội ngũ KTV HTM theo từng thiết bị "
    "y tế. Module đảm bảo chỉ KTV đã được đào tạo và xác nhận năng lực mới được phép thực hiện "
    "PM, CM hoặc Calibration trên thiết bị — tuân thủ ISO 13485 §6.2."
)
h3("Chức năng chính")
bullet("Lập và quản lý Training Program theo device category/model")
bullet("Ghi nhận kết quả đào tạo từng KTV (Training Session → Participant)")
bullet("Competency Gap Report: xác định KTV thiếu năng lực theo device assignment")
bullet("Competency Alert: cảnh báo năng lực sắp hết hiệu lực")
bullet("Authorized Technician list: danh sách KTV được phép vận hành từng thiết bị")

doc.add_paragraph()

# ── IMM-08 ───────────────────────────────────────────────────────────────────
h2("3.7 IMM-08 — Bảo trì định kỳ (Preventive Maintenance)")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Preventive Maintenance — Bảo trì định kỳ theo lịch"),
        ("Khối", "C — Khối 3"),
        ("Wave", "Wave 1 — Live"),
        ("Actor chính", "Workshop Manager, KTV HTM, PTP Khối 2"),
        ("DocType chính", "PM Schedule, PM Checklist Template, PM Work Order (submittable), PM Checklist Result, PM Task Log (immutable)"),
        ("Workflow", "IMM-08 PM Workflow — 7 states"),
        ("Naming series", "WO-PM-.YY.-.#####"),
        ("Số endpoints", "23 REST endpoints"),
    ]
)
h3("Mục đích")
para(
    "IMM-08 tự động hóa toàn bộ vòng đời PM: từ tạo PM Schedule khi commissioning, scheduler hàng ngày "
    "sinh PM Work Order đúng hạn, KTV điền checklist chuẩn hóa, đến cập nhật ngày PM kế tiếp và "
    "phát sinh CM Work Order khi phát hiện lỗi trong PM (Fail-Major)."
)
h3("Workflow states")
add_table(
    ["Trạng thái", "Mô tả"],
    [
        ("Scheduled", "WO được tạo bởi scheduler — chờ phân công"),
        ("Assigned", "Đã phân công KTV"),
        ("In Progress", "KTV đang thực hiện PM"),
        ("Completed — Pass", "PM hoàn thành, tất cả items Pass hoặc Fail-Minor"),
        ("Halted — Major Failure", "PM phát hiện lỗi nghiêm trọng → auto trigger IMM-09 CM"),
        ("Overdue", "Quá due_date chưa hoàn thành"),
        ("Cancelled", "Hủy WO"),
    ]
)
h3("Business Rules quan trọng")
bullet("BR-08-01: Scheduler daily tạo WO 7 ngày trước due_date, không tạo duplicate cho cùng schedule")
bullet("BR-08-02: Fail-Major → asset_status → Out of Service, auto-create CM WO trong IMM-09")
bullet("BR-08-03: PM Task Log là immutable sau submit — không xoá, không sửa")
bullet("BR-08-04: Checklist phải điền đủ 100% items trước khi submit")
bullet("BR-08-05: Workshop Manager mới được submit/cancel WO")
h3("Scheduler jobs")
add_table(
    ["Job", "Tần suất", "Mô tả"],
    [
        ("generate_due_pm_work_orders", "Daily", "Tạo PM WO cho các schedule đến hạn trong 7 ngày"),
        ("mark_pm_overdue", "Daily", "Đánh dấu Overdue các WO quá due_date"),
        ("send_pm_escalation_email", "Daily", "Gửi email leo thang cho Workshop Manager"),
        ("rollup_pm_compliance_kpi", "Monthly", "Tính PM Compliance Rate hàng tháng"),
    ]
)
h3("KPI mục tiêu")
add_table(
    ["KPI", "Baseline", "Target"],
    [
        ("PM Compliance Rate", "~60% (Excel manual)", "≥ 90%"),
        ("Số WO Overdue", "N/A", "≤ 5% tổng WO"),
        ("Avg Days Late", "N/A", "≤ 2 ngày"),
        ("PM Task Log Coverage", "0%", "100%"),
    ]
)

doc.add_paragraph()

# ── IMM-09 ───────────────────────────────────────────────────────────────────
h2("3.8 IMM-09 — Sửa chữa (Corrective Maintenance)")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Corrective Maintenance / Repair"),
        ("Khối", "C — Khối 3"),
        ("Wave", "Wave 1 — Live (reference module)"),
        ("Actor chính", "Workshop Manager, KTV HTM, Kho vật tư, Trưởng khoa phòng, PTP Khối 2"),
        ("DocType chính", "Asset Repair (submittable), Spare Parts Used, Repair Checklist, Firmware Change Request (submittable)"),
        ("Workflow", "IMM-09 Repair Workflow — 9 states"),
        ("Naming series", "WO-RP-.YY.-.#####"),
        ("Số endpoints", "12 REST endpoints"),
    ]
)
h3("Mục đích")
para(
    "IMM-09 chuẩn hóa toàn bộ vòng đời sửa chữa corrective: tiếp nhận từ Incident Report (IMM-12) "
    "hoặc PM Halted (IMM-08), phân công KTV, chẩn đoán, xuất vật tư có chứng từ kế toán, "
    "nghiệm thu 100% Pass, đo MTTR theo SLA risk class."
)
h3("Workflow states (9 trạng thái)")
add_table(
    ["Trạng thái", "Mô tả"],
    [
        ("Open", "WO mới tạo — chờ phân công"),
        ("Assigned", "Đã phân công KTV"),
        ("Diagnosing", "Đang chẩn đoán nguyên nhân"),
        ("Waiting Parts", "Chờ vật tư thay thế"),
        ("In Repair", "Đang tiến hành sửa chữa"),
        ("QA Verification", "Đang nghiệm thu chức năng — checklist 100% Pass"),
        ("Completed", "Sửa chữa thành công — asset_status → Active"),
        ("Cannot Repair", "Không thể sửa chữa → asset_status → Out of Service → trigger IMM-13/14"),
        ("Cancelled", "Hủy WO"),
    ]
)
h3("Business Rules quan trọng")
bullet("BR-09-01: Repair Checklist phải Pass 100% trước khi submit Completed")
bullet("BR-09-02: Spare Parts Used phải có stock_entry_ref từ ERPNext Stock")
bullet("BR-09-03: MTTR được tính từ created_date đến completion_date")
bullet("BR-09-04: SLA breach log tự động ghi khi MTTR > SLA threshold")
bullet("BR-09-05: Cannot Repair → bắt buộc tạo Incident Report nếu chưa có")
bullet("BR-09-06: Asset Repair record là immutable sau submit")
bullet("BR-09-07: Firmware change phải có FCR (Firmware Change Request) riêng biệt")
h3("SLA Matrix")
add_table(
    ["Risk Class", "Priority", "MTTR Target"],
    [
        ("Class III (Critical)", "Emergency", "≤ 4 giờ"),
        ("Class III (Critical)", "High", "≤ 24 giờ"),
        ("Class II (High)", "Normal", "≤ 72 giờ"),
        ("Class I (Low)", "Normal", "≤ 5 ngày"),
    ]
)
h3("Tích hợp")
add_table(
    ["Module", "Chiều", "Mô tả"],
    [
        ("IMM-08", "INPUT", "PM Halted–Major Failure auto-create CM WO"),
        ("IMM-12", "INPUT", "Incident Report → nguồn tạo CM WO"),
        ("IMM-11", "OUTPUT", "Post-repair → trigger Calibration nếu cần"),
        ("IMM-12", "OUTPUT", "Cannot Repair + Repeat failure → CAPA trigger"),
        ("AC Spare Part", "INPUT/OUTPUT", "Xuất/trả vật tư với stock_entry_ref"),
    ]
)
h3("KPI mục tiêu")
add_table(
    ["KPI", "Target"],
    [
        ("MTTR Class III Emergency", "≤ 4 giờ"),
        ("SLA Compliance Rate", "≥ 90%"),
        ("First-Time Fix Rate", "≥ 80%"),
        ("Repeat Failure Rate", "≤ 5%"),
    ]
)

doc.add_paragraph()

# ── IMM-11 ───────────────────────────────────────────────────────────────────
h2("3.9 IMM-11 — Hiệu chuẩn (Calibration)")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Hiệu năng và Hiệu chuẩn thiết bị đo lường y tế"),
        ("Khối", "C — Khối 3"),
        ("Wave", "Wave 1 — Live"),
        ("Actor chính", "IMM Workshop Lead, IMM Technician, IMM QA Officer"),
        ("DocType chính", "IMM Asset Calibration (submittable), IMM Calibration Schedule, IMM Calibration Measurement"),
        ("Workflow", "IMM-11 Calibration Workflow"),
        ("Naming series", "CAL-.YY.-.#####"),
        ("Số endpoints", "18 REST endpoints"),
    ]
)
h3("Mục đích")
para(
    "IMM-11 đảm bảo thiết bị đo lường y tế luôn trong dung sai cho phép bằng cách tự động lập lịch "
    "hiệu chuẩn, track bàn giao lab ISO/IEC 17025, tính kết quả Pass/Fail theo tolerance, và "
    "kích hoạt CAPA bắt buộc khi Out-of-Tolerance — không thiết bị nào fail vẫn dùng trên bệnh nhân."
)
h3("Loại hiệu chuẩn")
add_table(
    ["Loại", "Mô tả"],
    [
        ("External", "Bàn giao lab ISO/IEC 17025 bên ngoài — có certificate traceability"),
        ("In-House", "KTV nội bộ thực hiện với reference standard được chuẩn hóa"),
    ]
)
h3("Workflow states")
add_table(
    ["Trạng thái", "Mô tả"],
    [
        ("Scheduled", "Lịch hiệu chuẩn đến hạn — WO được tạo"),
        ("In Progress", "Đang bàn giao lab / KTV đang thực hiện"),
        ("Measurement Entered", "Đã nhập số liệu đo kiểm"),
        ("Pass", "Kết quả trong dung sai — asset tiếp tục active"),
        ("Fail — OOT", "Out-of-Tolerance → asset_status → Clinical Hold, CAPA bắt buộc"),
        ("Cancelled", "Hủy lịch (asset decommissioned)"),
    ]
)
h3("Business Rules")
bullet("BR-11-01: Tạo Calibration Schedule tự động khi IMM-04 submit (commissioning)")
bullet("BR-11-02: OOT → bắt buộc tạo CAPA Record và Lookback cho các asset cùng model")
bullet("BR-11-03: External calibration phải có ISO/IEC 17025 certificate upload")
bullet("BR-11-04: Pass/Fail tính tự động từ measured_value vs reference_value × tolerance_%")
bullet("BR-11-05: Asset bị OOT phải qua IMM-09 hoặc IMM-12 trước khi recalibrate")
bullet("BR-11-06: Suspend Calibration Schedule khi asset Decommissioned")
h3("KPI mục tiêu")
add_table(
    ["KPI", "Target"],
    [
        ("Calibration Compliance Rate", "≥ 95%"),
        ("Out-of-Tolerance (OOT) Rate", "< 5%"),
        ("Certificate Coverage", "100% External cals có certificate"),
        ("CAPA Triggered on OOT", "100%"),
    ]
)

doc.add_paragraph()

# ── IMM-12 ───────────────────────────────────────────────────────────────────
h2("3.10 IMM-12 — Sự cố & CAPA (Incident, RCA, CAPA)")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Sự cố thiết bị y tế, Phân tích nguyên nhân gốc rễ & Hành động khắc phục"),
        ("Khối", "C — Khối 3"),
        ("Wave", "Wave 1 — Live"),
        ("Actor chính", "Reporting User, IMM Workshop Lead, IMM QA Officer, IMM Operations Manager"),
        ("DocType chính", "Incident Report (submittable), IMM RCA Record, IMM CAPA Record, IMM RCA Five Why Step, IMM RCA Related Incident"),
        ("Workflow", "IMM-12 Incident Workflow, IMM-12 RCA Workflow"),
        ("Naming series", "IR-.YY.-.#####"),
        ("Số endpoints", "14 REST endpoints"),
    ]
)
h3("Mục đích")
para(
    "IMM-12 giải quyết vấn đề sự cố thiết bị y tế không được theo dõi hệ thống, dẫn đến lặp lại "
    "(chronic failure) mà không phát hiện. Module tự động phân loại mức độ, kích hoạt RCA bắt buộc "
    "với Major/Critical, tạo CAPA và phát hiện sự cố mãn tính qua scheduler hàng ngày."
)
h3("Phân loại sự cố")
add_table(
    ["Severity", "Tiêu chí", "Action bắt buộc"],
    [
        ("Minor", "Ảnh hưởng tối thiểu, không ảnh hưởng lâm sàng", "Ghi nhận, assign KTV"),
        ("Major", "Gián đoạn dịch vụ lâm sàng, nguy cơ trung bình", "RCA bắt buộc"),
        ("Critical", "Nguy hiểm bệnh nhân trực tiếp, thiệt hại tài sản lớn", "RCA + CAPA + Escalation tức thì"),
    ]
)
h3("Workflow Incident")
para("Open → Assigned → Under Investigation → Resolved → (nếu Major/Critical) → RCA Created → CAPA Linked → Closed")
h3("RCA Methods hỗ trợ")
bullet("5-Why Analysis (IMM RCA Five Why Step child table)")
bullet("Fishbone/Ishikawa (6M categories trong notes)")
bullet("Related Incident linking để phát hiện chronic failure pattern")
h3("Scheduler: Chronic Failure Detection")
para(
    "Scheduler hàng ngày phân tích: nếu cùng asset_ref có ≥ 3 incident trong 30 ngày → "
    "tự động escalate lên IMM Workshop Lead và IMM QA Officer để đánh giá decommission."
)
h3("KPI mục tiêu")
add_table(
    ["KPI", "Target"],
    [
        ("Incident Resolution Rate trong SLA", "≥ 85%"),
        ("RCA Completion Rate cho Major/Critical", "100%"),
        ("CAPA On-Time Closure Rate", "≥ 90%"),
        ("Chronic Failure Detection Lead Time", "≤ 30 ngày"),
    ]
)

doc.add_paragraph()

# ── IMM-15 ───────────────────────────────────────────────────────────────────
h2("3.11 IMM-15 — Spare Parts Inventory Tracking")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Theo dõi tồn kho phụ tùng y tế"),
        ("Khối", "C — Khối 3"),
        ("Wave", "Wave 3 — Planned"),
        ("Actor chính", "IMM Workshop Lead, Kho vật tư, IMM Technician"),
        ("DocType chính", "IMM Spare Part Forecast, IMM Spare Allocation, IMM Spare Allocation Item, IMM Stock Cycle Count, IMM Critical Spare Watchlist, IMM Spare Batch"),
        ("Workflow", "IMM-15 Allocation Workflow, IMM-15 Cycle Count Workflow"),
    ]
)
h3("Mục đích")
para(
    "IMM-15 quản lý tồn kho phụ tùng thiết bị y tế: dự báo nhu cầu dựa trên lịch sử CM (IMM-09), "
    "theo dõi stock level theo critical spare watchlist, và thực hiện cycle count định kỳ để "
    "đảm bảo spare parts sẵn sàng cho sửa chữa khẩn cấp."
)

doc.add_paragraph()

# ── IMM-16 ───────────────────────────────────────────────────────────────────
h2("3.12 IMM-16 — Giám sát Tuân thủ & CAPA")
add_table(
    ["Thuộc tính", "Giá trị"],
    [
        ("Tên đầy đủ", "Giám sát tuân thủ, Internal Audit & Management Review"),
        ("Khối", "C — Khối 3"),
        ("Wave", "Wave 3 — Planned"),
        ("Actor chính", "IMM QA Officer, IMM Department Head, IMM Operations Manager"),
        ("DocType chính", "IMM Internal Audit, IMM Compliance Rule, IMM Compliance Scorecard, IMM Compliance Finding, IMM Management Review"),
        ("Workflow", "IMM-16 CAPA Workflow, IMM-16 Finding Workflow, IMM-16 Internal Audit, IMM-16 MR Workflow"),
    ]
)
h3("Mục đích")
para(
    "IMM-16 chuẩn hóa quy trình audit nội bộ, theo dõi compliance theo bộ quy tắc định nghĩa "
    "và management review — đảm bảo hệ thống HTM của bệnh viện luôn tuân thủ NĐ 98/2021 và "
    "các tiêu chuẩn ISO áp dụng."
)

doc.add_paragraph()

# ── Wave 3 modules ────────────────────────────────────────────────────────────
h2("3.13 Nhóm module Wave 3 — IMM-07, 10, 13, 14, 17")
add_table(
    ["Module", "Tên", "Mục đích tóm tắt"],
    [
        ("IMM-07", "Theo dõi hiệu suất", "Dashboard KPI thiết bị — uptime, MTTR, OEE, cost-per-unit theo asset category"),
        ("IMM-10", "Hậu kiểm và tuân thủ", "Post-market surveillance: theo dõi field safety notice, recall từ nhà sản xuất"),
        ("IMM-13", "Ngừng sử dụng & điều chuyển", "Workflow ngừng sử dụng + điều chuyển thiết bị giữa khoa/cơ sở — có phê duyệt đa cấp"),
        ("IMM-14", "Giải nhiệm thiết bị", "Quy trình formal retirement: lập biên bản, xử lý tài sản, kết thúc lifecycle"),
        ("IMM-17", "Phân tích dự đoán", "Predictive maintenance cockpit: ML model dự báo hỏng hóc dựa trên sensor + maintenance history"),
    ]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 4 — DATA MODEL TỔNG THỂ
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Data model tổng thể")

h2("4.1 Master Data")
para(
    "Master data là dữ liệu nền tảng ít thay đổi, dùng chung cho toàn hệ thống:"
)
add_table(
    ["DocType", "Prefix", "Module", "Mô tả"],
    [
        ("AC Asset", "—", "IMM-00", "Registry thiết bị vật lý — trục trung tâm"),
        ("AC Asset Category", "—", "IMM-00", "Phân loại thiết bị theo IFU/risk class"),
        ("AC Supplier", "AC-SUP-", "IMM-00", "Nhà cung cấp (vendor, lab, NCC phụ tùng)"),
        ("AC Location", "AC-LOC-", "IMM-00", "Vị trí vật lý phân cấp"),
        ("AC Department", "AC-DEPT-", "IMM-00", "Khoa/phòng bệnh viện"),
        ("IMM Device Model", "IMM-MDL-", "IMM-00", "Catalog model thiết bị — PM interval, cal interval, risk class"),
        ("IMM SLA Policy", "—", "IMM-00", "Chính sách SLA theo risk_class × priority"),
        ("AC Spare Part", "—", "IMM-00/15", "Danh mục phụ tùng"),
        ("PM Checklist Template", "—", "IMM-08", "Template checklist PM theo category × pm_type"),
        ("IMM Compliance Rule", "—", "IMM-16", "Bộ quy tắc tuân thủ"),
    ]
)

h2("4.2 Operational Data")
para("Operational data là bản ghi giao dịch nghiệp vụ — submittable, có workflow:")
add_table(
    ["DocType", "Naming", "Module", "Mô tả"],
    [
        ("Asset Commissioning", "ACC-.YY.-.MM.-#####", "IMM-04", "Phiếu lắp đặt & commissioning"),
        ("Asset Document", "—", "IMM-05", "Hồ sơ tài liệu thiết bị"),
        ("IMM Training Program", "—", "IMM-06", "Chương trình đào tạo"),
        ("IMM Training Session", "—", "IMM-06", "Buổi đào tạo thực tế"),
        ("PM Schedule", "—", "IMM-08", "Lịch PM cho từng thiết bị"),
        ("PM Work Order", "WO-PM-.YY.-#####", "IMM-08", "Lệnh công việc PM"),
        ("PM Task Log", "—", "IMM-08", "Log immutable sau submit PM"),
        ("Asset Repair", "WO-RP-.YY.-#####", "IMM-09", "Lệnh sửa chữa CM"),
        ("Firmware Change Request", "—", "IMM-09", "Yêu cầu thay đổi firmware"),
        ("IMM Calibration Schedule", "—", "IMM-11", "Lịch hiệu chuẩn"),
        ("IMM Asset Calibration", "CAL-.YY.-#####", "IMM-11", "Phiếu hiệu chuẩn"),
        ("Incident Report", "IR-.YY.-#####", "IMM-12", "Báo cáo sự cố"),
        ("IMM RCA Record", "—", "IMM-12", "Phân tích nguyên nhân gốc rễ"),
        ("IMM Needs Request", "—", "IMM-01", "Yêu cầu nhu cầu thiết bị"),
        ("IMM Tech Spec", "—", "IMM-02", "Thông số kỹ thuật"),
        ("IMM Vendor Evaluation", "—", "IMM-03", "Đánh giá nhà cung cấp"),
        ("IMM Procurement Decision", "—", "IMM-03", "Quyết định mua sắm"),
    ]
)

h2("4.3 Governance Data")
para("Governance data là bản ghi quản trị — thường immutable hoặc append-only:")
add_table(
    ["DocType", "Module", "Mô tả"],
    [
        ("IMM Audit Trail", "IMM-00", "Audit trail SHA-256 chain — tuyệt đối immutable"),
        ("Asset Lifecycle Event", "IMM-00", "Sự kiện vòng đời — immutable sau tạo"),
        ("IMM CAPA Record", "IMM-00/16", "Corrective and Preventive Action"),
        ("Asset QA Non Conformance", "IMM-04", "Non-conformance trong commissioning"),
        ("IMM Calibration Measurement", "IMM-11", "Kết quả đo kiểm hiệu chuẩn"),
        ("IMM Compliance Finding", "IMM-16", "Phát hiện không tuân thủ"),
        ("IMM Internal Audit", "IMM-16", "Biên bản audit nội bộ"),
        ("IMM Management Review", "IMM-16", "Biên bản management review"),
        ("AC Asset Downtime Log", "IMM-00", "Log downtime thiết bị"),
        ("Expiry Alert Log", "IMM-05", "Log cảnh báo hết hạn tài liệu"),
    ]
)

h2("4.4 Danh mục DocType đầy đủ")
para(
    "Hệ thống AssetCore có tổng cộng 80+ DocType, được prefix theo domain:"
)
add_table(
    ["Prefix", "Domain", "Ví dụ"],
    [
        ("AC ", "Core entity (foundation)", "AC Asset, AC Supplier, AC Location, AC Department, AC Spare Part"),
        ("IMM ", "Governance/operational (modules)", "IMM Audit Trail, IMM Device Model, IMM CAPA Record, IMM Calibration Schedule"),
        ("Asset ", "Operational records (Wave 1)", "Asset Commissioning, Asset Repair, Asset Document, Asset Lifecycle Event"),
        ("PM ", "Preventive Maintenance", "PM Schedule, PM Work Order, PM Task Log, PM Checklist Template"),
        ("Incident ", "Incident management", "Incident Report"),
        ("Service ", "Contracts", "Service Contract, Service Contract Asset"),
    ]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 5 — TÍCH HỢP HỆ THỐNG NGOÀI
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Tích hợp hệ thống ngoài")

h2("5.1 REST API / OpenAPI")
para(
    "AssetCore expose tất cả chức năng qua REST API chuẩn Frappe Whitelist. "
    "Response envelope thống nhất: "
)
para('  { "success": true/false, "data": {...} | [...], "error": "...", "code": "..." }', italic=True)
add_table(
    ["Module", "Số endpoints", "Nhóm chức năng"],
    [
        ("IMM-00", "70+", "Asset CRUD, SLA, Audit Trail, Lifecycle Event, CAPA"),
        ("IMM-04", "17", "Commissioning, Gate status, Approval flow, QR mint"),
        ("IMM-05", "14", "Document CRUD, Completeness, Expiry alert"),
        ("IMM-08", "23", "PM Schedule, Work Order, Checklist, KPI dashboard"),
        ("IMM-09", "12", "Repair WO, Spare parts, Checklist, MTTR"),
        ("IMM-11", "18", "Calibration schedule, Measurement, Pass/Fail"),
        ("IMM-12", "14", "Incident, RCA, CAPA, Chronic detection"),
        ("IMM-01/02/03", "~30", "Needs, Tech Spec, Vendor Eval, Procurement"),
    ]
)

h2("5.2 FHIR Integration (Roadmap)")
para(
    "AssetCore được thiết kế FHIR-ready để tích hợp với HIS/EMR. "
    "AC Asset ánh xạ sang FHIR Device resource. Asset Lifecycle Event ánh xạ sang "
    "FHIR DeviceUseStatement. Incident Report ánh xạ sang FHIR AdverseEvent."
)
add_table(
    ["AssetCore Entity", "FHIR Resource", "Trạng thái"],
    [
        ("AC Asset", "Device", "Roadmap"),
        ("Asset Lifecycle Event", "DeviceUseStatement", "Roadmap"),
        ("Incident Report", "AdverseEvent", "Roadmap"),
        ("IMM Calibration", "Observation (device)", "Roadmap"),
    ]
)

h2("5.3 Hệ thống bệnh viện")
add_table(
    ["Hệ thống", "Chiều tích hợp", "Dữ liệu trao đổi"],
    [
        ("HIS (Hospital Information System)", "Bidirectional", "Danh sách thiết bị, downtime alert, maintenance schedule"),
        ("ERP/Tài chính", "OUTPUT", "Asset record → khấu hao, inventory valuation"),
        ("LIS/RIS/PACS", "INPUT", "Incident từ modality (imaging equipment failure)"),
        ("GMDN (Global Medical Device Nomenclature)", "LOOKUP", "Phân loại thiết bị theo mã GMDN quốc tế"),
    ]
)

h2("5.4 Authentication & Security")
para(
    "AssetCore dùng Frappe built-in session authentication. "
    "Permission được kiểm soát ở 2 lớp: Frappe DocPerm (document-level) và "
    "API permission check trong service layer (role-based). "
    "Tất cả API endpoint đều có @frappe.whitelist() decorator với explicit role check."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 6 — WORKFLOW ENGINE & SLA
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Workflow engine & SLA")

h2("6.1 Work Order Engine")
para(
    "Work Order là engine trung tâm của AssetCore — mọi action trên thiết bị đều phải "
    "đi qua một Work Order tương ứng:"
)
add_table(
    ["Loại Work Order", "DocType", "Module", "Trigger"],
    [
        ("Commissioning", "Asset Commissioning", "IMM-04", "Manual — từ PO"),
        ("Preventive Maintenance", "PM Work Order", "IMM-08", "Scheduler daily (auto)"),
        ("Corrective Maintenance", "Asset Repair", "IMM-09", "Incident / PM Halted (auto/manual)"),
        ("Calibration", "IMM Asset Calibration", "IMM-11", "Scheduler / Post-repair (auto/manual)"),
        ("Incident Report", "Incident Report", "IMM-12", "Manual / auto từ IMM-08/09/11"),
    ]
)

h2("6.2 SLA Policy Engine")
para(
    "SLA Policy được giải quyết theo 3 chiều: risk_class × priority × asset_category."
)
add_table(
    ["Trường", "Mô tả"],
    [
        ("risk_class", "Class I (Low) / Class II (High) / Class III (Critical) — theo IEC 62366/NĐ 98"),
        ("priority", "Emergency / High / Normal / Low"),
        ("asset_category", "Imaging / Life Support / Laboratory / General / Infra"),
        ("response_hours", "Thời gian phản hồi tối đa (giờ)"),
        ("resolution_hours", "Thời gian giải quyết tối đa (giờ) — MTTR target"),
    ]
)
para(
    "Function resolve_sla_policy(asset_ref, priority) trả về SLA Policy record "
    "phù hợp nhất theo priority cascade."
)

h2("6.3 Asset Lifecycle Event")
para(
    "Mỗi state mutation của AC Asset đều sinh Asset Lifecycle Event row. "
    "Event là immutable — không xoá, không sửa sau khi tạo."
)
add_table(
    ["Field", "Mô tả"],
    [
        ("asset_ref", "Link đến AC Asset"),
        ("event_type", "commissioned / pm_completed / repaired / calibrated / failure_reported / retired / ..."),
        ("from_status", "Lifecycle status trước khi chuyển"),
        ("to_status", "Lifecycle status sau khi chuyển"),
        ("actor", "User thực hiện action"),
        ("timestamp", "Datetime UTC — immutable"),
        ("root_record", "Doctype + docname của Work Order/Commissioning nguồn"),
    ]
)

h2("6.4 Audit Trail (SHA-256 Chain)")
para(
    "IMM Audit Trail là bản ghi bất biến cho mọi state mutation nghiệp vụ. "
    "Mỗi record chứa SHA-256 hash của nội dung và hash của record trước — "
    "tạo thành chain không thể tamper mà không phá vỡ toàn chuỗi."
)
add_table(
    ["Field", "Mô tả"],
    [
        ("event_type", "Loại sự kiện (submit_commissioning, repair_completed, calibration_fail, ...)"),
        ("document_type", "DocType nguồn"),
        ("document_name", "Docname nguồn"),
        ("actor", "User thực hiện"),
        ("timestamp", "Datetime UTC"),
        ("payload_json", "Nội dung đầy đủ tại thời điểm action"),
        ("hash_current", "SHA-256(event_type + timestamp + actor + payload + hash_previous)"),
        ("hash_previous", "Hash của record ngay trước — tạo chain"),
    ]
)
para("Verify chain: function verify_chain() tính lại hash toàn chuỗi và xác nhận không có tamper.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 7 — BẢO MẬT & PHÂN QUYỀN
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Bảo mật & Phân quyền")

h2("7.1 Vai trò hệ thống (8 roles)")
add_table(
    ["Role", "Mô tả", "Access level"],
    [
        ("IMM System Admin", "Quản trị viên hệ thống HTM", "Full CRUD + config + fixtures"),
        ("IMM Department Head", "Trưởng phòng HTM / BGĐ kỹ thuật", "Read all + Approve decisions"),
        ("IMM Operations Manager", "Quản lý vận hành TBYT", "CRUD operational records"),
        ("IMM Workshop Lead", "Trưởng xưởng kỹ thuật", "CRUD WO + Submit/Cancel"),
        ("IMM Technician", "Kỹ thuật viên HTM", "Read/Update WO gán cho mình"),
        ("IMM QA Officer", "Nhân viên QA/QLCL", "Review + Approve QA actions + CAPA"),
        ("IMM Document Officer", "Nhân viên lưu trữ tài liệu", "Read-only + Export"),
        ("IMM Reporting User", "User báo cáo sự cố (khoa phòng)", "Create Incident Report only"),
    ]
)

h2("7.2 Nguyên tắc phân quyền")
bullet("Principle of Least Privilege: mỗi role chỉ có quyền tối thiểu cần thiết")
bullet("KHÔNG gán System Manager cho operational roles")
bullet("Vendor Engineer: isolated — chỉ xem phiếu commissioning được gán, không xem AC Asset của cơ sở")
bullet("DocPerm: mỗi DocType có permission matrix riêng trong JSON fixture")
bullet("API permission: service layer check frappe.has_permission() trước mọi mutation")
bullet("Audit trail: không role nào có quyền xoá IMM Audit Trail hoặc Asset Lifecycle Event")

h2("7.3 Security controls")
add_table(
    ["Control", "Implement"],
    [
        ("Authentication", "Frappe session + CSRF token"),
        ("Authorization", "Role-based DocPerm + service layer role check"),
        ("Audit trail", "SHA-256 chained — tamper detection"),
        ("SQL injection", "Frappe ORM parameterized queries — không raw SQL"),
        ("Mass assignment", "@frappe.whitelist() explicit payload validation"),
        ("Vendor isolation", "User Permission filter — vendor chỉ xem record của mình"),
        ("Immutable records", "Lifecycle Event + Audit Trail: no_copy, no delete permission"),
    ]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 8 — YÊU CẦU PHI CHỨC NĂNG
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Yêu cầu phi chức năng")

add_table(
    ["Nhóm", "Yêu cầu", "Target"],
    [
        ("Performance", "API response time (p95)", "< 500ms cho list queries, < 200ms cho get single"),
        ("Performance", "FE Time to Interactive", "< 3 giây trên kết nối 10Mbps"),
        ("Performance", "Scheduler job execution", "< 60 giây cho daily batch"),
        ("Availability", "Uptime", "≥ 99.5% trong giờ hành chính"),
        ("Scalability", "Số asset đồng thời", "≥ 10,000 assets không degradation"),
        ("Data integrity", "Audit trail completeness", "100% — 0 state mutation không có audit"),
        ("Data integrity", "Serial number uniqueness", "100% — 0 duplicate"),
        ("Compliance", "NĐ 98/2021 traceability", "100% thiết bị có lifecycle traceable từ commissioning"),
        ("Compliance", "ISO 13485 audit readiness", "Export đầy đủ hồ sơ cho audit trong < 5 phút"),
        ("Security", "SHA-256 chain integrity", "verify_chain() pass 100% sau mọi audit trail insert"),
        ("Maintainability", "Test coverage BE service layer", "≥ 80%"),
        ("Localization", "Ngôn ngữ", "Tiếng Việt UI, tiếng Anh field names/code"),
    ]
)

h2("8.1 Technology Stack")
add_table(
    ["Lớp", "Công nghệ", "Phiên bản"],
    [
        ("Backend Framework", "Frappe Framework", "v15"),
        ("Backend Language", "Python", "3.11+"),
        ("Database", "MariaDB", "10.6+"),
        ("Frontend Framework", "Vue 3 + TypeScript", "Vue 3.4+"),
        ("State Management", "Pinia", "2.x"),
        ("HTTP Client", "TanStack Query + axios", "—"),
        ("UI Library", "TailwindCSS", "3.x"),
        ("Testing BE", "pytest + frappe.tests", "—"),
        ("Testing FE", "Vitest + Playwright", "—"),
        ("CI/CD", "GitHub Actions", "—"),
    ]
)

h2("8.2 Deployment Architecture")
add_table(
    ["Component", "Mô tả"],
    [
        ("Application server", "Frappe/Gunicorn + Nginx reverse proxy"),
        ("Database server", "MariaDB (tách instance khỏi app server ở production)"),
        ("Queue worker", "Redis + Frappe Worker (background jobs + scheduler)"),
        ("File storage", "Frappe file system (local) / S3-compatible (cloud)"),
        ("FE assets", "Vite build → dist/ → Nginx static serve"),
        ("SSL", "Let's Encrypt / hospital cert"),
    ]
)

doc.add_paragraph()
note(
    "Tài liệu này được sinh tự động từ codebase AssetCore tại "
    "/home/miyano/frappe-bench/apps/assetcore — "
    "phiên bản 1.0.0 ngày 2026-05-13. "
    "Để cập nhật, chạy lại script gen_spec.py sau khi có thay đổi docs/imm-*/."
)

doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT):,} bytes")
