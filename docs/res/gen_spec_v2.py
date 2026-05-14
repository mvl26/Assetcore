#!/usr/bin/env python3
"""Tạo lại AssetCore_Technical_Specification.docx theo yêu cầu mới."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "/home/miyano/frappe-bench/apps/assetcore/docs/res/AssetCore_Technical_Specification.docx"

doc = Document()

# ── Cấu hình trang ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)

# ── Styles cơ bản ──────────────────────────────────────────────────────────
styles = doc.styles

normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(13)

# Heading 1 — tiêu đề phần
h1 = styles["Heading 1"]
h1.font.name = "Times New Roman"
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# Heading 2 — tên module
h2 = styles["Heading 2"]
h2.font.name = "Times New Roman"
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

# Heading 3 — tên chức năng
h3 = styles["Heading 3"]
h3.font.name = "Times New Roman"
h3.font.size = Pt(13)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def para(text, bold=False, italic=False, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold  = bold
    run.italic = italic
    return p


def heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
    return h


def feature(name, desc):
    """In đậm tên chức năng rồi mô tả liền bên dưới."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(name)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)

    p2 = doc.add_paragraph(desc)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p2.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
    p2.paragraph_format.left_indent = Cm(0.5)
    return p, p2


def caption(text):
    """Chú thích hình minh họa — in đậm + nghiêng, căn giữa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    return p


# ════════════════════════════════════════════════════════════════════════════
# TRANG BÌA
# ════════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("ASSETCORE")
run.font.name = "Times New Roman"
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run("Hệ thống Quản lý Vòng đời Thiết bị Y tế")
run2.font.name = "Times New Roman"
run2.font.size = Pt(18)
run2.bold = True

doc.add_paragraph()

p_spec = doc.add_paragraph()
p_spec.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p_spec.add_run("Đặc tả Chức năng Phần mềm")
run3.font.name = "Times New Roman"
run3.font.size = Pt(16)

doc.add_paragraph()
doc.add_paragraph()

p_ver = doc.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p_ver.add_run("Phiên bản 1.0  •  2026")
run4.font.name = "Times New Roman"
run4.font.size = Pt(13)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PHẦN 1 — GIỚI THIỆU PHẦN MỀM
# ════════════════════════════════════════════════════════════════════════════
heading("Phần 1. Giới thiệu phần mềm", 1)

para(
    "AssetCore là phần mềm quản lý toàn bộ vòng đời thiết bị y tế tại bệnh viện, "
    "từ giai đoạn xác định nhu cầu và lên kế hoạch mua sắm, qua lắp đặt và đưa vào vận hành, "
    "đến bảo trì định kỳ, sửa chữa, hiệu chuẩn, và cuối cùng là thanh lý hoặc điều chuyển. "
    "Phần mềm thay thế toàn bộ quy trình giấy tờ và bảng tính Excel hiện tại bằng một hệ thống "
    "điện tử thống nhất, có phân quyền theo vai trò, có lịch sử thao tác không thể xóa, "
    "và có cảnh báo tự động theo thời hạn."
)

para(
    "AssetCore dành cho các đơn vị y tế — bệnh viện công và tư — có phòng kỹ thuật thiết bị y tế "
    "hoặc nhóm bảo trì nội bộ. Người dùng bao gồm kỹ thuật viên thiết bị y tế, kỹ sư y sinh, "
    "cán bộ quản lý kho phụ tùng, cán bộ quản lý chất lượng, ban giám đốc phụ trách kỹ thuật, "
    "và nhân viên kinh tế tài sản."
)

para(
    "Phần mềm xây dựng theo chuẩn WHO HTM (Health Technology Management) và phù hợp với "
    "yêu cầu của Nghị định 98/2021/NĐ-CP về quản lý thiết bị y tế tại Việt Nam. "
    "Mọi thao tác tác động đến trạng thái thiết bị đều được ghi nhận tự động với thông tin "
    "người thực hiện, thời gian, và nội dung thay đổi."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PHẦN 2 — DANH SÁCH MODULE VÀ CHỨC NĂNG
# ════════════════════════════════════════════════════════════════════════════
heading("Phần 2. Danh sách module và chức năng", 1)

# ── 2.1 Module Dữ liệu Tham chiếu ──────────────────────────────────────────
heading("2.1. Module Dữ liệu Tham chiếu", 2)

para(
    "Module này cung cấp toàn bộ danh mục dùng chung cho các module khác trong hệ thống: "
    "danh mục thiết bị, khoa phòng, địa điểm, nhà cung cấp, phân loại nguy cơ và chính sách bảo trì."
)

feature(
    "Quản lý danh mục thiết bị (Model thiết bị)",
    "Người quản trị tạo hồ sơ model thiết bị với đầy đủ thông tin: tên thiết bị, hãng sản xuất, "
    "mã GMDN, phân loại nguy cơ (Class A/B/C/D theo NĐ98), chu kỳ bảo trì khuyến nghị, "
    "và các tài liệu kỹ thuật đính kèm. Model thiết bị là nền tảng để tạo hồ sơ thiết bị thực tế "
    "và thiết lập lịch bảo trì tự động."
)
caption("Hình 2.1.1: Trang danh sách model thiết bị")

feature(
    "Quản lý khoa phòng",
    "Hệ thống cho phép tạo và quản lý danh mục khoa phòng trong bệnh viện. Người quản trị vào Module "
    "Dữ liệu Tham chiếu, chọn mục Khoa Phòng, nhấn \"Tạo mới\" để nhập tên và mã tự động sinh. "
    "Sau khi lưu, khoa phòng được sử dụng làm dữ liệu tham chiếu trong toàn hệ thống "
    "(gắn với thiết bị, lịch bảo trì, phiếu sửa chữa)."
)
caption("Hình 2.1.2: Màn hình tạo mới khoa phòng")

feature(
    "Quản lý địa điểm",
    "Cho phép xây dựng cây phân cấp địa điểm theo tòa nhà, tầng, phòng. Khi thiết bị được lắp đặt "
    "hoặc điều chuyển, địa điểm mới được cập nhật tự động vào hồ sơ thiết bị, đảm bảo tồn kho luôn "
    "phản ánh đúng vị trí thực tế."
)

feature(
    "Quản lý nhà cung cấp",
    "Lưu trữ thông tin nhà cung cấp thiết bị y tế: tên, địa chỉ, mã số thuế, hồ sơ pháp lý (giấy "
    "phép kinh doanh, chứng chỉ phân phối hãng) kèm ngày hết hạn. Hệ thống cảnh báo tự động khi "
    "hồ sơ pháp lý sắp hết hạn, giúp bộ phận mua sắm chủ động yêu cầu cập nhật trước khi ký hợp đồng mới."
)

feature(
    "Quản lý chính sách bảo trì",
    "Người quản lý xưởng thiết lập chính sách bảo trì theo từng loại thiết bị: chu kỳ bảo trì (tháng), "
    "thời gian thực hiện tối đa (giờ), và mức ưu tiên theo phân loại nguy cơ. "
    "Chính sách này được áp dụng tự động khi tạo lịch bảo trì định kỳ cho thiết bị mới."
)

# ── 2.2 Module Lập kế hoạch Nhu cầu ─────────────────────────────────────────
heading("2.2. Module Lập kế hoạch Nhu cầu Thiết bị", 2)

para(
    "Module này hỗ trợ toàn bộ quy trình từ tiếp nhận đề xuất mua thiết bị từ các khoa phòng, "
    "đánh giá ưu tiên, lập dự toán, đến tổng hợp kế hoạch mua sắm trình ban giám đốc phê duyệt."
)

feature(
    "Tạo đề xuất nhu cầu thiết bị",
    "Trưởng khoa hoặc kỹ thuật viên vào màn hình Đề xuất Nhu cầu, điền thông tin thiết bị cần mua "
    "(tên, số lượng, lý do lâm sàng, ước tính kinh phí) rồi gửi đề xuất lên phòng kỹ thuật. "
    "Mỗi đề xuất được cấp mã tự động và có trạng thái theo dõi từ khi tạo đến khi được duyệt hoặc từ chối."
)
caption("Hình 2.2.1: Màn hình tạo đề xuất nhu cầu thiết bị")

feature(
    "Chấm điểm ưu tiên đa tiêu chí",
    "Cán bộ kế hoạch chấm điểm mỗi đề xuất theo sáu tiêu chí: tác động lâm sàng, mức độ nguy cơ "
    "khi thiếu thiết bị, tỷ lệ sử dụng hiện tại, tín hiệu hư hỏng của thiết bị cũ, mức độ tuân thủ "
    "quy định pháp lý, và khả năng nguồn vốn. Tổng điểm giúp sắp xếp thứ tự ưu tiên khách quan "
    "và có thể kiểm toán sau này."
)

feature(
    "Lập dự toán tổng chi phí sở hữu",
    "Với mỗi thiết bị trong danh sách ưu tiên, hệ thống hỗ trợ tính tổng chi phí sở hữu gồm: "
    "chi phí mua ban đầu (vốn đầu tư), chi phí vận hành hàng năm (vật tư tiêu hao, điện, nhân công), "
    "chi phí bảo trì và hiệu chuẩn trong 5 năm, và chi phí thanh lý ước tính. "
    "Kết quả dự toán được lưu vào hồ sơ đề xuất và theo sang giai đoạn phê duyệt."
)

feature(
    "Tổng hợp kế hoạch mua sắm",
    "Sau khi chấm điểm và có dự toán, cán bộ kế hoạch tổng hợp thành Kế hoạch Mua sắm cho cả năm. "
    "Ban giám đốc xem xét và phê duyệt trên hệ thống; khi kế hoạch được phê duyệt, hệ thống tự động "
    "mở quy trình soạn thảo thông số kỹ thuật."
)
caption("Hình 2.2.2: Trang tổng hợp kế hoạch mua sắm theo năm")

# ── 2.3 Module Mua sắm ───────────────────────────────────────────────────────
heading("2.3. Module Mua sắm Thiết bị", 2)

para(
    "Module này quản lý quy trình từ soạn thảo thông số kỹ thuật, đánh giá và lựa chọn nhà cung cấp, "
    "đến phát hành quyết định mua sắm và đơn đặt hàng."
)

feature(
    "Soạn thảo thông số kỹ thuật",
    "Kỹ sư y sinh soạn thông số kỹ thuật chi tiết cho từng thiết bị cần mua: yêu cầu về tính năng, "
    "tiêu chuẩn điện, kết nối hạ tầng, và khả năng tương thích với hệ thống thông tin bệnh viện. "
    "Tài liệu được kiểm soát phiên bản và phải được khóa ở trạng thái phê duyệt trước khi "
    "tiến hành đánh giá nhà cung cấp."
)

feature(
    "Đánh giá và lựa chọn nhà cung cấp",
    "Cán bộ mua sắm chấm điểm tối thiểu ba nhà cung cấp theo năm nhóm tiêu chí: kỹ thuật, "
    "thương mại, tài chính, tuân thủ pháp lý, và dịch vụ sau bán hàng. Kết quả chấm điểm "
    "được lưu vào hồ sơ đấu thầu và không thể sửa sau khi đã nộp. Hệ thống tự kiểm tra "
    "hình thức mua sắm phù hợp theo ngưỡng giá trị và loại thiết bị."
)
caption("Hình 2.3.1: Bảng chấm điểm đánh giá nhà cung cấp")

feature(
    "Quản lý danh sách nhà cung cấp được duyệt",
    "Hệ thống duy trì danh sách nhà cung cấp được phê duyệt theo từng loại thiết bị. "
    "Nhà cung cấp chỉ có thể tham gia đánh giá khi còn trong danh sách hợp lệ và hồ sơ pháp lý "
    "chưa hết hạn. Workshop Head định kỳ tái đánh giá nhà cung cấp và cập nhật điểm phản hồi "
    "từ các lần lắp đặt và sửa chữa trước đó."
)

feature(
    "Phát hành đơn đặt hàng",
    "Sau khi quyết định mua sắm được phê duyệt, hệ thống tạo đơn đặt hàng liên kết với quyết định "
    "đó. Đơn đặt hàng ghi nhận nhà cung cấp, thiết bị, số lượng, giá trị, điều khoản bảo hành, "
    "và thời hạn giao hàng. Tất cả đều có mã tham chiếu để đối soát khi thiết bị về đến kho."
)

# ── 2.4 Module Lắp đặt ───────────────────────────────────────────────────────
heading("2.4. Module Lắp đặt Thiết bị", 2)

para(
    "Module này kiểm soát toàn bộ quy trình từ khi thiết bị về kho đến khi chính thức đưa vào "
    "sử dụng lâm sàng: kiểm tra hồ sơ, lắp đặt vật lý, gán mã định danh, đo kiểm an toàn điện, "
    "và phê duyệt đưa vào vận hành."
)

feature(
    "Tạo phiếu lắp đặt từ đơn đặt hàng",
    "Khi thiết bị về kho, kỹ thuật viên tạo phiếu lắp đặt bằng cách chọn đơn đặt hàng tương ứng. "
    "Hệ thống tự điền thông tin thiết bị, nhà cung cấp, và khoa phòng dự kiến lắp đặt. "
    "Mỗi phiếu có mã tự động theo định dạng chuẩn và được theo dõi theo trạng thái qua toàn bộ "
    "quy trình nghiệm thu."
)
caption("Hình 2.4.1: Trang danh sách phiếu lắp đặt thiết bị")

feature(
    "Kiểm tra hồ sơ tiếp nhận",
    "Kỹ thuật viên đánh dấu danh sách các tài liệu đi kèm thiết bị (chứng nhận xuất xứ, chứng nhận "
    "hợp quy, hướng dẫn sử dụng, phiếu bảo hành). Với thiết bị có nguy cơ cao (Class C, D), "
    "hệ thống yêu cầu giấy phép bức xạ hoặc giấy phép nhập khẩu trước khi cho phép tiếp tục."
)

feature(
    "Gán mã định danh nội bộ và mã QR",
    "Sau khi lắp đặt, kỹ sư y sinh nhập số serial của nhà sản xuất và hệ thống tự sinh mã định danh "
    "nội bộ theo định dạng chuẩn của bệnh viện. Mã QR được tạo tự động, có thể in ra để dán lên "
    "thiết bị. Số serial được kiểm tra trùng lặp trong toàn hệ thống."
)

feature(
    "Đo kiểm an toàn điện và kiểm tra lâm sàng ban đầu",
    "Kỹ thuật viên điền kết quả đo kiểm an toàn điện (dòng rò, điện trở bảo vệ) vào bảng kiểm chuẩn "
    "hóa trên hệ thống. Cán bộ QA xem xét kết quả; nếu đạt, hồ sơ chuyển lên phó trưởng khối "
    "ký duyệt lần cuối. Nếu không đạt, phiếu trở về trạng thái chờ xử lý với lý do ghi nhận rõ ràng."
)

feature(
    "Phê duyệt và tạo hồ sơ thiết bị chính thức",
    "Khi phiếu lắp đặt được phê duyệt hoàn toàn, hệ thống tự động tạo hồ sơ thiết bị chính thức "
    "trong registry với trạng thái vận hành, đồng thời mở lịch bảo trì định kỳ đầu tiên và "
    "chuyển bộ hồ sơ tài liệu vào Module Hồ sơ Thiết bị. Thiết bị chỉ được phép sử dụng trên "
    "bệnh nhân sau bước này."
)
caption("Hình 2.4.2: Phiếu lắp đặt — bảng kiểm an toàn điện và ký duyệt")

# ── 2.5 Module Hồ sơ Thiết bị ────────────────────────────────────────────────
heading("2.5. Module Hồ sơ Thiết bị", 2)

para(
    "Module này là kho lưu trữ tập trung tất cả tài liệu kỹ thuật và pháp lý gắn với từng thiết bị "
    "trong suốt vòng đời: từ hướng dẫn sử dụng, sơ đồ mạch điện, chứng chỉ hiệu chuẩn, "
    "đến giấy phép bức xạ và hợp đồng bảo hành."
)

feature(
    "Lưu trữ và phân loại tài liệu theo thiết bị",
    "Kỹ thuật viên tải tài liệu lên và gán nhãn loại: hướng dẫn sử dụng, sơ đồ kỹ thuật, "
    "chứng chỉ an toàn, hồ sơ pháp lý, hoặc chứng chỉ hiệu chuẩn. Tài liệu được liên kết "
    "với hồ sơ thiết bị cụ thể và có thể tìm kiếm nhanh theo mã thiết bị hoặc loại tài liệu. "
    "Mỗi lần tải lên tạo một phiên bản mới, phiên bản cũ vẫn được giữ lại để kiểm toán."
)

feature(
    "Phê duyệt tài liệu pháp lý",
    "Tài liệu pháp lý quan trọng (giấy phép, chứng chỉ an toàn) phải qua luồng duyệt: kỹ thuật viên "
    "tải lên, tổ HC-QLCL xem xét và phê duyệt hoặc từ chối với lý do. Chỉ tài liệu đã được duyệt "
    "mới được tính là hợp lệ trong các cổng kiểm tra tuân thủ."
)

feature(
    "Cảnh báo hết hạn tài liệu",
    "Với mỗi tài liệu có ngày hết hạn (giấy phép bức xạ, hợp đồng bảo hành, chứng chỉ hiệu chuẩn), "
    "hệ thống gửi cảnh báo tự động trước 90, 60, và 30 ngày đến người phụ trách. "
    "Dashboard hiển thị danh sách tài liệu sắp hết hạn theo thiết bị và khoa phòng."
)
caption("Hình 2.5.1: Trang hồ sơ tài liệu của một thiết bị")

feature(
    "Lưu trữ tự động từ các quy trình khác",
    "Sau mỗi lần lắp đặt, hiệu chuẩn, hoặc sửa chữa hoàn thành, hệ thống tự động sao chép "
    "biên bản nghiệm thu hoặc chứng chỉ vào kho hồ sơ của thiết bị tương ứng. "
    "Người dùng không cần tải thủ công — toàn bộ lịch sử tài liệu được xây dựng tự động theo vòng đời."
)

# ── 2.6 Module Đào tạo và Quản lý Năng lực ───────────────────────────────────
heading("2.6. Module Đào tạo và Quản lý Năng lực", 2)

para(
    "Module này quản lý chứng nhận đào tạo và năng lực vận hành thiết bị y tế theo NĐ98/2021 "
    "và ISO 13485, đảm bảo chỉ người được đào tạo mới được phân công thực hiện Work Order."
)

feature(
    "Quản lý chương trình đào tạo",
    "Tổ HC-QLCL thiết kế chương trình đào tạo cho từng loại thiết bị y tế: nội dung học, "
    "thời lượng, phương pháp đánh giá và điểm đạt tối thiểu. Chương trình được gắn với "
    "phân loại thiết bị để tự động áp dụng khi có thiết bị mới đưa vào vận hành."
)

feature(
    "Tổ chức buổi đào tạo và ghi nhận kết quả",
    "Workshop Lead tạo buổi đào tạo, mời người tham dự, và nhập điểm sau khi kết thúc. "
    "Hệ thống tự động tính kết quả đạt hoặc không đạt theo ngưỡng của chương trình. "
    "Người tham dự có thể xem chứng nhận của mình trên hệ thống ngay sau khi kết quả được xác nhận."
)

feature(
    "Cấp và theo dõi chứng nhận năng lực",
    "Sau khi đạt đào tạo, hệ thống tự động cấp chứng nhận năng lực cho người dùng theo thiết bị "
    "và loại công việc (vận hành, bảo trì, hiệu chuẩn). Chứng nhận có ngày hết hạn; khi sắp hết hạn, "
    "hệ thống nhắc nhở tái đào tạo."
)

feature(
    "Kiểm tra năng lực trước khi phân công công việc",
    "Khi Workshop Lead phân công kỹ thuật viên thực hiện bảo trì hoặc sửa chữa, hệ thống tự kiểm tra "
    "xem người đó có chứng nhận hợp lệ cho thiết bị đó không. Nếu chứng nhận hết hạn hoặc chưa có, "
    "hệ thống cảnh báo và yêu cầu xác nhận trước khi tiếp tục."
)
caption("Hình 2.6.1: Trang danh sách chứng nhận năng lực của kỹ thuật viên")

# ── 2.7 Module Quản lý Thiết bị Vận hành ──────────────────────────────────────
heading("2.7. Module Quản lý Thiết bị Vận hành", 2)

para(
    "Module này theo dõi trạng thái, vị trí, và hiệu suất hoạt động của toàn bộ thiết bị đang "
    "trong giai đoạn vận hành, cung cấp bức tranh tổng thể cho ban quản lý."
)

feature(
    "Theo dõi trạng thái thiết bị",
    "Mỗi thiết bị có trạng thái rõ ràng: đang hoạt động, đang bảo trì, đang sửa chữa, tạm ngừng, "
    "hoặc hết khả năng sửa. Trạng thái được cập nhật tự động khi có Work Order mở hoặc đóng. "
    "Người dùng xem danh sách thiết bị với bộ lọc theo khoa, trạng thái, và loại thiết bị."
)
caption("Hình 2.7.1: Trang danh sách thiết bị vận hành — lọc theo trạng thái")

feature(
    "Theo dõi điều chuyển thiết bị nội viện",
    "Khi thiết bị được di chuyển từ khoa này sang khoa khác, kỹ thuật viên tạo phiếu điều chuyển "
    "với lý do và chữ ký xác nhận của trưởng khoa nhận. Hồ sơ thiết bị tự động cập nhật "
    "địa điểm mới, đảm bảo tồn kho luôn chính xác."
)

feature(
    "Bảng điều khiển hiệu suất vận hành",
    "Dashboard tổng hợp các chỉ số vận hành theo từng thiết bị, loại thiết bị, và khoa phòng: "
    "tỷ lệ sẵn sàng hoạt động, thời gian ngừng máy, số lần sự cố, và thời gian sửa chữa trung bình. "
    "Ban giám đốc truy cập bảng điều khiển này để nắm tình trạng thiết bị theo thời gian thực."
)
caption("Hình 2.7.2: Bảng điều khiển hiệu suất vận hành thiết bị")

feature(
    "Phát hiện tín hiệu cần thay thế thiết bị",
    "Hệ thống phân tích xu hướng hiệu suất của từng thiết bị. Khi một thiết bị có số lần hỏng tăng "
    "liên tục hoặc thời gian sửa chữa kéo dài vượt ngưỡng, hệ thống tự động phát tín hiệu cảnh báo "
    "cho Workshop Head và cán bộ kế hoạch để xem xét đề xuất thay thế."
)

# ── 2.8 Module Bảo trì Định kỳ ────────────────────────────────────────────────
heading("2.8. Module Bảo trì Định kỳ", 2)

para(
    "Module này tự động hóa toàn bộ vòng đời bảo trì định kỳ: từ tạo lịch khi thiết bị được "
    "lắp đặt, đến sinh Work Order đúng hạn, hướng dẫn kỹ thuật viên thực hiện bảo trì "
    "theo bảng kiểm chuẩn hóa, và theo dõi tỷ lệ tuân thủ bảo trì."
)

feature(
    "Tạo và quản lý lịch bảo trì",
    "Ngay khi thiết bị được nghiệm thu lắp đặt, hệ thống tự động tạo lịch bảo trì định kỳ "
    "dựa trên chính sách bảo trì của loại thiết bị đó. Workshop Head có thể điều chỉnh chu kỳ "
    "và ngày bảo trì tiếp theo. Lịch bảo trì của toàn bộ thiết bị hiển thị dạng lịch (calendar) "
    "để dễ lên kế hoạch nhân lực."
)
caption("Hình 2.8.1: Lịch bảo trì định kỳ dạng calendar")

feature(
    "Tự động tạo Work Order bảo trì",
    "Hệ thống kiểm tra lịch bảo trì mỗi ngày và tự động tạo Work Order khi đến hạn. "
    "Work Order được gán cho kỹ thuật viên có năng lực phù hợp, kèm bảng kiểm đã soạn sẵn "
    "cho loại thiết bị đó. Nếu quá hạn mà chưa có người nhận, hệ thống cảnh báo Workshop Head."
)

feature(
    "Thực hiện bảo trì theo bảng kiểm chuẩn hóa",
    "Kỹ thuật viên mở Work Order được phân công, điền kết quả từng hạng mục trong bảng kiểm: "
    "đạt (Pass), hỏng nhỏ (Fail-Minor), hỏng lớn (Fail-Major), hoặc không áp dụng (N/A). "
    "Sau khi hoàn thành tất cả hạng mục, kỹ thuật viên xác nhận hoàn thành. Nếu có hạng mục "
    "Fail-Major, hệ thống tự động đặt thiết bị vào trạng thái chờ sửa chữa và tạo Work Order sửa chữa."
)
caption("Hình 2.8.2: Work Order bảo trì — bảng kiểm kết quả từng hạng mục")

feature(
    "Theo dõi tỷ lệ tuân thủ bảo trì",
    "Dashboard hiển thị tỷ lệ bảo trì hoàn thành đúng hạn (mục tiêu tối thiểu 90%), số Work Order "
    "quá hạn, và xu hướng tuân thủ theo tháng. Workshop Head sử dụng thông tin này để báo cáo "
    "cho ban giám đốc và cải thiện kế hoạch nhân lực bảo trì."
)

# ── 2.9 Module Sửa chữa Thiết bị ──────────────────────────────────────────────
heading("2.9. Module Sửa chữa Thiết bị", 2)

para(
    "Module này chuẩn hóa toàn bộ quy trình sửa chữa thiết bị: từ tiếp nhận yêu cầu, "
    "chẩn đoán, xuất vật tư, đến nghiệm thu và đóng hồ sơ sửa chữa."
)

feature(
    "Tiếp nhận yêu cầu sửa chữa",
    "Work Order sửa chữa được tạo tự động từ báo cáo sự cố (Module Quản lý Sự cố) hoặc từ "
    "kết quả bảo trì phát hiện hỏng lớn (Module Bảo trì Định kỳ). Workshop Head cũng có thể "
    "tạo thủ công khi nhận yêu cầu trực tiếp. Mỗi Work Order có mức ưu tiên (thường, khẩn, "
    "khẩn cấp) theo phân loại nguy cơ của thiết bị."
)
caption("Hình 2.9.1: Trang danh sách Work Order sửa chữa — lọc theo mức ưu tiên")

feature(
    "Chẩn đoán và ghi nhận nguyên nhân hỏng",
    "Kỹ thuật viên được phân công nhập mô tả hiện tượng hỏng, chẩn đoán nguyên nhân gốc, "
    "và lập danh sách vật tư cần dùng. Thông tin này được lưu vĩnh viễn vào hồ sơ Work Order "
    "và dùng cho phân tích xu hướng hỏng hóc về sau."
)

feature(
    "Xuất vật tư phụ tùng theo Work Order",
    "Kỹ thuật viên yêu cầu xuất vật tư từ kho; thủ kho xác nhận và ghi nhận xuất kho gắn với "
    "mã Work Order. Không có xuất vật tư ngoài Work Order trừ trường hợp khẩn cấp có ghi lý do. "
    "Hồ sơ xuất kho tự động cập nhật tồn kho."
)

feature(
    "Nghiệm thu và đóng Work Order",
    "Sau khi sửa xong, kỹ thuật viên điền kết quả kiểm tra nghiệm thu. Trưởng khoa nơi sử dụng "
    "xác nhận thiết bị đã hoạt động tốt. Work Order được đóng và thiết bị tự động trở lại "
    "trạng thái hoạt động. Thời gian sửa chữa được tính toán tự động để phân tích MTTR."
)
caption("Hình 2.9.2: Màn hình nghiệm thu và đóng Work Order sửa chữa")

feature(
    "Theo dõi thỏa thuận mức dịch vụ (SLA) sửa chữa",
    "Hệ thống theo dõi thời gian xử lý của từng Work Order so với ngưỡng SLA quy định theo "
    "phân loại nguy cơ thiết bị. Khi sắp vi phạm SLA, hệ thống gửi cảnh báo leo thang đến "
    "Workshop Head. Dashboard hiển thị tỷ lệ tuân thủ SLA theo tháng và theo loại thiết bị."
)

# ── 2.10 Module Hiệu chuẩn Thiết bị ──────────────────────────────────────────
heading("2.10. Module Hiệu chuẩn Thiết bị", 2)

para(
    "Module này quản lý lịch hiệu chuẩn định kỳ, theo dõi quá trình bàn giao thiết bị cho "
    "phòng lab đo lường, ghi nhận kết quả so sánh với chuẩn đo lường quốc gia, "
    "và kích hoạt hành động khắc phục khi thiết bị vượt dung sai."
)

feature(
    "Lập và quản lý lịch hiệu chuẩn",
    "Lịch hiệu chuẩn được tạo tự động khi thiết bị được nghiệm thu lắp đặt, dựa trên chu kỳ "
    "khuyến nghị của nhà sản xuất (thường 12 tháng). Hệ thống tạo Work Order hiệu chuẩn trước "
    "30 ngày so với ngày đến hạn, đủ thời gian liên hệ phòng lab và chuẩn bị bàn giao thiết bị."
)

feature(
    "Theo dõi bàn giao cho phòng lab và nhận kết quả",
    "Kỹ thuật viên ghi nhận ngày bàn giao thiết bị cho phòng lab hiệu chuẩn (nội bộ hoặc bên ngoài). "
    "Khi nhận lại, kỹ thuật viên tải chứng chỉ hiệu chuẩn lên và nhập số liệu đo được. "
    "Hệ thống tự động tính kết quả đạt hoặc không đạt dựa trên dung sai cho phép của từng chỉ số."
)
caption("Hình 2.10.1: Trang nhập kết quả hiệu chuẩn và tải chứng chỉ")

feature(
    "Xử lý khi thiết bị không đạt hiệu chuẩn",
    "Khi thiết bị không đạt hiệu chuẩn, hệ thống tự động đặt thiết bị vào trạng thái tạm ngừng "
    "và mở phiếu CAPA bắt buộc. Đồng thời kích hoạt quy trình nhìn lại (Lookback) để xác định "
    "xem trong thời gian thiết bị vượt dung sai có kết quả xét nghiệm hoặc đo lường nào trên "
    "bệnh nhân cần được xem xét lại không. Thiết bị chỉ được đưa lại vào sử dụng sau khi "
    "có chứng chỉ hiệu chuẩn đạt."
)

feature(
    "Báo cáo tuân thủ hiệu chuẩn",
    "Dashboard hiệu chuẩn hiển thị tỷ lệ thiết bị đã hiệu chuẩn đúng hạn, danh sách thiết bị "
    "sắp đến hạn trong 30–60 ngày tới, và lịch sử kết quả hiệu chuẩn của từng thiết bị. "
    "Báo cáo này phục vụ kiểm toán nội bộ và báo cáo cho cơ quan quản lý nhà nước."
)

# ── 2.11 Module Quản lý Sự cố ─────────────────────────────────────────────────
heading("2.11. Module Quản lý Sự cố", 2)

para(
    "Module này tiếp nhận, phân loại và theo dõi các sự cố thiết bị y tế từ khi báo cáo đến khi "
    "xác định nguyên nhân gốc, thực hiện hành động khắc phục, và đóng hồ sơ sự cố."
)

feature(
    "Báo cáo sự cố nhanh",
    "Điều dưỡng hoặc kỹ thuật viên khoa phòng có thể báo sự cố trực tiếp từ thiết bị di động. "
    "Thông tin cần điền tối thiểu gồm: thiết bị nào, hiện tượng gì, thời điểm xảy ra, và mức độ "
    "ảnh hưởng tức thời đến người bệnh. Sau khi gửi, hệ thống tự động thông báo cho Workshop Head "
    "và tạo Work Order sửa chữa nếu thiết bị cần ngừng ngay."
)
caption("Hình 2.11.1: Màn hình báo cáo sự cố nhanh")

feature(
    "Phân loại mức độ nghiêm trọng",
    "Workshop Head phân loại sự cố thành ba mức: nhỏ (Minor), lớn (Major), hoặc nghiêm trọng "
    "(Critical) theo tiêu chí chuẩn về tác động đến người bệnh và tính an toàn lâm sàng. "
    "Sự cố mức Major và Critical bắt buộc phải có phân tích nguyên nhân gốc (RCA)."
)

feature(
    "Phân tích nguyên nhân gốc (RCA)",
    "Với sự cố mức Major hoặc Critical, hệ thống mở phiếu RCA và hướng dẫn người thực hiện "
    "điền theo phương pháp có cấu trúc (5-Why hoặc Fishbone). Kết quả RCA được lưu vào hồ sơ "
    "sự cố và trở thành căn cứ cho phiếu CAPA."
)

feature(
    "Tạo và theo dõi phiếu CAPA",
    "Từ hồ sơ sự cố, cán bộ QA tạo phiếu hành động khắc phục và phòng ngừa (CAPA) với hạn hoàn "
    "thành cụ thể và người chịu trách nhiệm. Hệ thống theo dõi tiến độ và cảnh báo khi sắp quá hạn. "
    "Phiếu CAPA chỉ được đóng khi hành động khắc phục đã được xác nhận hiệu quả."
)
caption("Hình 2.11.2: Trang theo dõi tiến độ phiếu CAPA")

feature(
    "Phát hiện sự cố mãn tính",
    "Hệ thống tự động phân tích dữ liệu mỗi ngày để phát hiện thiết bị có số lần sự cố tăng bất "
    "thường trong 30, 60, hoặc 90 ngày gần nhất. Khi phát hiện mẫu này, hệ thống tạo cảnh báo "
    "cho Workshop Head và cán bộ QA để điều tra nguyên nhân hệ thống."
)

# ── 2.12 Module Quản lý Phụ tùng Vật tư ──────────────────────────────────────
heading("2.12. Module Quản lý Phụ tùng Vật tư", 2)

para(
    "Module này quản lý tồn kho phụ tùng thiết bị y tế, cấp phát theo Work Order, "
    "theo dõi mức tồn tối thiểu, và hỗ trợ kiểm kê định kỳ."
)

feature(
    "Quản lý danh mục phụ tùng",
    "Kỹ thuật viên thiết lập danh mục phụ tùng với đầy đủ thông tin: mã phụ tùng, tên, "
    "nhà sản xuất, thiết bị tương thích, mức tồn kho tối thiểu, và thời gian đặt hàng tối thiểu. "
    "Phụ tùng Critical (phục vụ thiết bị nguy cơ cao) được đánh dấu và theo dõi riêng."
)

feature(
    "Cấp phát phụ tùng theo Work Order",
    "Kỹ thuật viên yêu cầu xuất phụ tùng từ kho bằng cách gắn yêu cầu với mã Work Order. "
    "Thủ kho xác nhận xuất kho; tồn kho tự động giảm và ghi nhận lịch sử tiêu hao theo "
    "từng Work Order và thiết bị được sửa."
)

feature(
    "Cảnh báo tồn kho phụ tùng quan trọng",
    "Hệ thống kiểm tra tồn kho mỗi ngày và gửi cảnh báo ngay khi bất kỳ phụ tùng Critical "
    "xuống dưới mức tồn tối thiểu. Workshop Head nhận thông báo tức thì để kịp thời đặt hàng, "
    "tránh tình trạng thiết bị phải dừng chờ phụ tùng."
)
caption("Hình 2.12.1: Trang tồn kho phụ tùng — danh sách cảnh báo dưới mức tối thiểu")

feature(
    "Kiểm kê định kỳ",
    "Workshop Head lên lịch kiểm kê; hệ thống tạo phiếu kiểm kê với danh sách phụ tùng cần đếm. "
    "Kỹ thuật viên nhập số lượng thực đếm; hệ thống tự động so sánh với sổ sách và hiển thị "
    "chênh lệch. Nếu chênh lệch vượt ngưỡng 5%, hệ thống yêu cầu mở phiếu điều tra nguyên nhân."
)

# ── 2.13 Module Thanh lý Thiết bị ─────────────────────────────────────────────
heading("2.13. Module Thanh lý Thiết bị", 2)

para(
    "Module này quản lý hai giai đoạn cuối vòng đời thiết bị: (1) quyết định ngừng sử dụng hoặc "
    "điều chuyển nội viện, và (2) đóng vĩnh viễn hồ sơ thiết bị kèm đối soát tài sản-kho-kế toán."
)

feature(
    "Đề xuất ngừng sử dụng thiết bị",
    "Kỹ thuật viên hoặc Trưởng khoa tạo đề xuất ngừng sử dụng với lý do bắt buộc: hết khả năng "
    "sửa chữa, vượt tuổi thọ thiết kế, thiếu phụ tùng không còn sản xuất, hoặc nhu cầu thay thế "
    "công nghệ. Hệ thống tự động tính điểm rủi ro còn lại dựa trên trạng thái thiết bị và lịch "
    "sử hỏng hóc."
)

feature(
    "Điều chuyển thiết bị trước khi thanh lý",
    "Nếu thiết bị còn sử dụng được ở khoa khác hoặc cơ sở vệ tinh, hệ thống hỗ trợ luồng "
    "điều chuyển nội viện với đầy đủ chữ ký điện tử của khoa nhận và khoa giao. "
    "Hồ sơ thiết bị cập nhật địa điểm và trạng thái mới ngay khi hoàn tất thủ tục."
)

feature(
    "Phê duyệt quyết định ngừng sử dụng",
    "Đề xuất ngừng sử dụng phải qua ba cấp duyệt: Workshop Head, tổ HC-QLCL, và ban giám đốc "
    "phụ trách kỹ thuật. Mỗi cấp có thể duyệt, từ chối (có lý do), hoặc yêu cầu bổ sung thông tin. "
    "Khi quyết định được duyệt, thiết bị chuyển sang giai đoạn đóng vòng đời."
)

feature(
    "Đóng vĩnh viễn vòng đời thiết bị",
    "Cán bộ tài sản thực hiện đối soát ba chiều: hồ sơ tài sản, phụ tùng còn tồn trong kho "
    "gắn với thiết bị này, và giá trị còn lại trên sổ kế toán. Sau đó nhập kết quả xử lý "
    "dữ liệu bệnh nhân (nếu có), đính kèm biên bản thanh lý, và xác nhận đóng. "
    "Thiết bị chuyển sang trạng thái đã thanh lý — không thể đảo ngược — và toàn bộ hồ sơ "
    "được lưu trữ vĩnh viễn để phục vụ kiểm toán."
)
caption("Hình 2.13.1: Phiếu đề xuất ngừng sử dụng thiết bị — luồng ký duyệt ba cấp")

# ── 2.14 Module Báo cáo và Thống kê ──────────────────────────────────────────
heading("2.14. Module Báo cáo và Thống kê", 2)

para(
    "Module này cung cấp hệ thống báo cáo và bảng điều khiển cho các cấp quản lý, "
    "từ Workshop Head đến ban giám đốc, với khả năng truy xuất đến hồ sơ nguồn."
)

feature(
    "Bảng điều khiển vận hành tổng thể",
    "Trang chủ hiển thị các chỉ số quan trọng theo thời gian thực: số thiết bị đang hoạt động "
    "và đang hỏng, tỷ lệ bảo trì đúng hạn, Work Order quá hạn, và thiết bị sắp đến hạn "
    "hiệu chuẩn. Người dùng nhấn vào bất kỳ chỉ số nào để đi thẳng đến danh sách chi tiết."
)
caption("Hình 2.14.1: Bảng điều khiển vận hành tổng thể")

feature(
    "Báo cáo tuân thủ bảo trì định kỳ",
    "Báo cáo liệt kê toàn bộ thiết bị theo tình trạng bảo trì: đã bảo trì đúng hạn, bảo trì trễ, "
    "và chưa đến hạn. Có thể lọc theo khoa phòng, loại thiết bị, và khoảng thời gian. "
    "Dùng để báo cáo cho ban giám đốc và cơ quan quản lý định kỳ."
)

feature(
    "Báo cáo sự cố và xu hướng hỏng hóc",
    "Thống kê số lượng sự cố theo mức độ nghiêm trọng, loại thiết bị, và khoa phòng trong khoảng "
    "thời gian được chọn. Biểu đồ xu hướng hiển thị sự thay đổi theo tháng, giúp phát hiện "
    "loại thiết bị hoặc khoa phòng có tần suất sự cố tăng bất thường."
)

feature(
    "Báo cáo hiệu suất tài sản (MTBF, MTTR)",
    "Mỗi thiết bị có trang hồ sơ hiệu suất với các chỉ số: thời gian hoạt động trung bình giữa "
    "hai lần hỏng (MTBF), thời gian sửa chữa trung bình (MTTR), và tỷ lệ sẵn sàng hoạt động. "
    "Dữ liệu được tổng hợp theo tháng và năm, hỗ trợ quyết định bảo trì và thay thế."
)

feature(
    "Xuất dữ liệu để báo cáo cơ quan quản lý",
    "Hệ thống hỗ trợ xuất báo cáo theo định dạng chuẩn để nộp Sở Y tế hoặc Bộ Y tế: "
    "danh sách thiết bị y tế đang quản lý, tình trạng đăng ký lưu hành, kết quả kiểm định "
    "định kỳ. Dữ liệu xuất ra ở định dạng Excel có thể tùy chỉnh theo yêu cầu báo cáo cụ thể."
)
caption("Hình 2.14.2: Trang báo cáo sự cố — biểu đồ xu hướng theo tháng")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PHẦN 3 — LUỒNG NGHIỆP VỤ XUYÊN SUỐT
# ════════════════════════════════════════════════════════════════════════════
heading("Phần 3. Luồng nghiệp vụ xuyên suốt", 1)

para(
    "Vòng đời một thiết bị y tế trong AssetCore bắt đầu từ khi bệnh viện nhận ra nhu cầu "
    "và kết thúc khi thiết bị được đóng vĩnh viễn khỏi danh mục hoạt động. "
    "Toàn bộ luồng này được kết nối liên tục — không có bước nào tách rời khỏi dữ liệu của "
    "bước trước và bước sau."
)

para(
    "Khi một khoa lâm sàng đề xuất mua thiết bị mới, yêu cầu đó vào hệ thống qua Module Lập "
    "kế hoạch Nhu cầu. Cán bộ kỹ thuật đánh giá, chấm điểm ưu tiên, lập dự toán chi phí toàn "
    "vòng đời, rồi tổng hợp thành kế hoạch mua sắm trình ban giám đốc. Khi kế hoạch được duyệt, "
    "Module Mua sắm tiếp nhận để soạn thông số kỹ thuật, đánh giá nhà cung cấp theo tiêu chí "
    "chuẩn hóa, và phát hành đơn đặt hàng có đầy đủ hồ sơ."
)

para(
    "Khi thiết bị về đến bệnh viện, Module Lắp đặt trở thành cổng kiểm tra bắt buộc. "
    "Không có phiếu lắp đặt được phê duyệt hoàn toàn thì thiết bị không được phép sử dụng trên "
    "bệnh nhân. Tại bước này, hệ thống kiểm tra hồ sơ pháp lý, gán mã định danh và mã QR nội bộ, "
    "ghi nhận kết quả đo kiểm an toàn điện, và thu thập chữ ký xác nhận từ tất cả các bên liên quan. "
    "Khi phiếu lắp đặt được phê duyệt, thiết bị xuất hiện trong registry với đầy đủ hồ sơ gốc "
    "và lịch bảo trì định kỳ được tạo tự động."
)

para(
    "Trong suốt giai đoạn vận hành, ba luồng nghiệp vụ chạy song song và liên kết chặt chẽ "
    "với nhau. Bảo trì định kỳ chạy theo lịch tự động: hệ thống tạo Work Order đúng hạn, "
    "kỹ thuật viên thực hiện theo bảng kiểm, và hệ thống cập nhật lịch bảo trì tiếp theo "
    "sau khi hoàn thành. Hiệu chuẩn cũng chạy theo chu kỳ tương tự: trước 30 ngày đến hạn, "
    "Work Order hiệu chuẩn được tạo để nhắc nhở liên hệ phòng lab. Sửa chữa xảy ra không theo "
    "lịch: khi thiết bị hỏng hoặc kết quả bảo trì phát hiện vấn đề lớn, Work Order sửa chữa "
    "tự động được tạo và theo dõi theo SLA."
)

para(
    "Mọi sự cố xảy ra với thiết bị đều được ghi nhận qua Module Quản lý Sự cố. Sự cố nghiêm "
    "trọng bắt buộc phân tích nguyên nhân gốc và mở phiếu CAPA. Nếu cùng một loại thiết bị "
    "xuất hiện sự cố lặp lại, hệ thống phát tín hiệu sự cố mãn tính để ban quản lý can thiệp. "
    "Khi thiết bị bị sửa chữa xong, nếu loại hỏng hóc đó liên quan đến độ chính xác đo lường, "
    "hệ thống tự động yêu cầu hiệu chuẩn lại trước khi đưa thiết bị trở lại sử dụng."
)

para(
    "Khi thiết bị tiếp cận cuối vòng đời — hỏng quá nhiều, hết phụ tùng thay thế, hoặc công "
    "nghệ lỗi thời — Module Thanh lý tiếp quản. Quy trình bắt đầu bằng đề xuất ngừng sử dụng "
    "có căn cứ rõ ràng và phân tích rủi ro còn lại. Nếu thiết bị còn có thể dùng ở nơi khác, "
    "luồng điều chuyển nội viện được thực hiện trước. Sau khi quyết định thanh lý được ký duyệt "
    "đầy đủ, cán bộ tài sản thực hiện đối soát ba chiều giữa hồ sơ kỹ thuật, tồn kho phụ tùng, "
    "và sổ kế toán, rồi chính thức đóng hồ sơ thiết bị với biên bản đầy đủ. "
    "Toàn bộ lịch sử của thiết bị từ ngày lắp đặt đến ngày thanh lý được lưu trữ vĩnh viễn "
    "trong hệ thống để phục vụ kiểm toán."
)

para(
    "Xuyên suốt toàn bộ vòng đời, Module Hồ sơ Thiết bị đóng vai trò kho lưu trữ tài liệu "
    "tập trung: mọi tài liệu phát sinh từ lắp đặt, hiệu chuẩn, sửa chữa đều tự động được "
    "lưu vào đây mà không cần nhập thủ công. Module Báo cáo và Thống kê theo dõi các chỉ số "
    "quan trọng và cung cấp thông tin cho ban giám đốc ra quyết định kịp thời. "
    "Module Quản lý Phụ tùng đảm bảo chuỗi cung ứng phụ tùng không gián đoạn, "
    "cảnh báo sớm khi phụ tùng quan trọng xuống mức nguy hiểm."
)

# ── Ghi chú cuối ──
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_footer.add_run("— Hết tài liệu —")
run_f.font.name = "Times New Roman"
run_f.font.size = Pt(11)
run_f.italic = True

doc.save(OUTPUT)
print(f"Da tao: {OUTPUT}")
